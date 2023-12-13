import os
import shutil
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from datetime import datetime, date
from dateutil.relativedelta import relativedelta
import pandas as pd
from docx import Document
import sqlite3
from tkinter import END
from tkinter import Toplevel
from openpyxl import load_workbook
import re
import subprocess
import sys
import openpyxl
import webbrowser
from pathlib import Path
from tkcalendar import Calendar
from tkinter.font import Font
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.styles import Alignment
from ttkthemes import ThemedTk
from openpyxl import Workbook
import math
from decimal import Decimal, ROUND_HALF_UP
from decimal import Decimal, InvalidOperation
from openpyxl.styles import PatternFill
from io import BytesIO
import threading
import io



from tkinter import simpledialog
from PIL import Image, ImageTk
from openpyxl_image_loader import SheetImageLoader

# Prototyping (make it work, then make it pretty.)

class DatabaseManager: #DB practice(use txt to store folder paths when program finished for faster reads.)

    def __init__(self, db_name='inventory_management.db'):
        self.conn = sqlite3.connect(db_name)
        self.cur = self.conn.cursor()
        self.setup_database()

    def setup_database(self):
        self.cur.execute('''
            CREATE TABLE IF NOT EXISTS folder_paths (
                Folder TEXT PRIMARY KEY,
                Path TEXT
            )
        ''')
        self.conn.commit()

    def save_folder_path(self, folder, path):
        self.cur.execute('''
            INSERT INTO folder_paths (Folder, Path) VALUES (?, ?)
            ON CONFLICT(Folder) DO UPDATE SET Path = excluded.Path;
        ''', (folder, path))
        self.conn.commit()

    def get_folder_path(self, folder_name):
        self.cur.execute('SELECT Path FROM folder_paths WHERE Folder = ?', (folder_name,))
        result = self.cur.fetchone()
        return result[0] if result else None

    def get_all_folders(self):
        self.cur.execute('SELECT Folder FROM folder_paths')
        return [row[0] for row in self.cur.fetchall()]

    def delete_all_folders(self):
        self.cur.execute('DELETE FROM folder_paths')
        self.conn.commit()
        
    def commit_changes(self):
        self.conn.commit()
        
    def __del__(self):
        if hasattr(self, 'conn'):
            self.conn.close()

class ExcelManager:

    def __init__(self, filepath=None, sheet_name=None):
        self.filepath = filepath
        self.sheet_name = sheet_name
        self.data_frame = None

    def load_data(self):
        if self.filepath and self.sheet_name:
            self.data_frame = pd.read_excel(self.filepath, sheet_name=self.sheet_name, engine='openpyxl')
            # Cast all columns to object dtype after loading data
            self.data_frame = self.data_frame.astype('object')

    def get_product_info(self, product_id):
        if self.data_frame is not None:
            # Convert both the product_id and the 'Product ID' column to lower case for comparison
            query_result = self.data_frame[self.data_frame['Product ID'].str.upper() == product_id.upper()]
            if not query_result.empty:
                return query_result.iloc[0].to_dict()
        return None

    def save_product_info(self, product_id, product_data):
        if self.filepath:
            try:
                #print(f"Loading workbook from {self.filepath}")
                workbook = load_workbook(self.filepath)
                #print(f"Accessing sheet {self.sheet_name}")
                sheet = workbook[self.sheet_name]

                # Start by finding the column index for product IDs
                product_id_col_index = self.get_column_index_by_header(sheet, 'Product ID')
                if not product_id_col_index:
                    #print("Product ID column not found")
                    return

                # Update product_data dictionary to convert boolean to YES/NO strings
                for key, value in product_data.items():
                    if isinstance(value, bool):
                        product_data[key] = 'YES' if value else 'NO'

                # Now iterate over the rows to find the matching product ID
                for row in sheet.iter_rows(min_col=product_id_col_index, max_col=product_id_col_index):
                    cell = row[0]
                    if cell.value and str(cell.value).strip().upper() == product_id.upper():
                        row_num = cell.row
                        for key, value in product_data.items():
                            col_index = self.get_column_index_by_header(sheet, key)
                            if col_index:
                                # Special handling for 'To Sell After' date
                                if key == 'To Sell After' and isinstance(value, datetime):
                                    value = value.strftime('%m/%d/%Y')  # Format the date
                                    sheet.cell(row=row_num, column=col_index, value=value)
                                elif key == 'Fair Market Value':
                                    # Convert value to float if it's not None or empty
                                    value = float(value) if value else 0
                                    # Set the cell value
                                    cell = sheet.cell(row=row_num, column=col_index, value=value)
                                    # Set the number format for currency
                                    cell.number_format = '"$"#,##0.00'
                                else:
                                    sheet.cell(row=row_num, column=col_index, value=value)
                        workbook.save(self.filepath)
                        break
                else:
                    #print(f"Product ID {product_id} not found in the sheet.")
                    pass
            except Exception as e:
                #print(f"Failed to save changes to Excel file: {e}")
                raise

    @staticmethod
    def get_column_index_by_header(sheet, header_name):
        """
        Gets the column index based on the header name.
        :param sheet: The sheet to search in.
        :param header_name: The header name to find.
        :return: The index of the column, or None if not found.
        """
        for col in sheet.iter_rows(min_row=1, max_row=1, values_only=True):
            if header_name in col:
                return col.index(header_name) + 1
        return None


class Application(tk.Frame):

    def __init__(self, master=None):
        super().__init__(master)
        self.db_manager = DatabaseManager()
        self.excel_manager = ExcelManager()
        self.edit_mode = False  # Add this line to initialize the edit_mode attribute
        self.inventory_folder = None
        self.sold_folder = None
        self.to_sell_folder = None
        self.pack(fill='both', expand=True)
        self.last_changed = None
        self.initial_discount_price = None  # Class attribute to store the initial discount price
        self.initial_percent_discount = None  # Class attribute to store the initial discount price
        self.initial_product_price_plus_ivu = ''  # Initialize the variable
        self.trigger_price_focus_out_flag = True
        self.running = True

        #self.trigger_save_flag = False # Can be used to save when pressing enter once while in Product Price (+IVU) entry.


        
        # Make sure you call this before combining and displaying folders
        self.Main_Window_Widgets() 
        
        # Now it's safe to load settings and combine folders since the list widget is created
        self.load_settings()
        self.combine_and_display_folders()
        
        # Call the methods associated with the settings buttons
        #self.update_links_in_excel()  # This corresponds to 'Autofill Excel Data(link, asin, tosellafter)'
        #self.update_folders_paths()   # This corresponds to 'Update folder names and paths'

    def close_application(self):
        self.running = False
        self.destroy()

    def load_settings(self):
        # Load settings
        try:
            with open("folders_paths.txt", "r") as file:
                lines = file.read().splitlines()
                self.inventory_folder = lines[0]
                self.sold_folder = lines[1]
                self.to_sell_folder = lines[2] if len(lines) > 2 else None
                # ... The rest of your settings loading code ...
        except FileNotFoundError:
            pass
        # Here you could handle the situation if the file is not found, like setting default paths or prompting the user.

    def save_settings(self):
        # This function is called after selecting the source and sold folders
        # Update the table with the new paths
        self.db_manager.cur.execute('''
            UPDATE folder_paths SET Path = ? WHERE Folder = 'Root Folder'
        ''', (self.inventory_folder,))
        self.db_manager.cur.execute('''
            UPDATE folder_paths SET Path = ? WHERE Folder = 'Sold'
        ''', (self.sold_folder,))
        self.db_manager.conn.commit()

    def check_and_update_product_list(self):
        if not self.search_entry.get():  # Check if the search entry is empty
            folder_count = len(next(os.walk(self.inventory_folder))[1])  # Count folders in the directory
            list_count = self.folder_list.size()  # Count items in the Listbox

            if folder_count != list_count:
                self.combine_and_display_folders()  # Update the list items with folder names

            # Schedule this method to be called again after 10000 milliseconds (10 seconds)
            #self.after(10000, self.check_and_update_product_list)

    def Main_Window_Widgets(self):
        
        self.top_frame = ttk.Frame(self)
        self.top_frame.pack(fill='x')

        self.settings_button = ttk.Button(self.top_frame, text='Settings', command=self.Settings_Window_Start)
        self.settings_button.pack(side='right')

        self.search_frame = ttk.Frame(self)
        self.search_frame.pack(fill='x')

        self.search_label = ttk.Label(self.search_frame, text="Enter product name here:")
        self.search_label.pack(anchor='w')

        self.search_entry = ttk.Entry(self.search_frame, width=30)  # Same width as the Listbox
        self.search_entry.pack(side='left', fill='x', anchor='w')
        self.search_entry.bind('<KeyRelease>', self.search)

        self.bottom_frame = ttk.Frame(self)
        self.bottom_frame.pack(fill='both', expand=True)

        self.list_outer_frame = ttk.Frame(self.bottom_frame)
        self.list_outer_frame.pack(side='left', fill='y')

        self.list_frame = ttk.Frame(self.list_outer_frame)
        self.list_frame.pack(side='left', fill='both', expand=True)

        self.folder_list = tk.Listbox(self.list_frame, width=30)
        self.folder_list.pack(side='left', fill='both', expand=False)
        self.folder_list.bind('<<ListboxSelect>>', self.display_product_details)

        self.list_scrollbar = ttk.Scrollbar(self.list_frame)
        self.list_scrollbar.pack(side='right', fill='y')
        self.folder_list.config(yscrollcommand=self.list_scrollbar.set)
        self.list_scrollbar.config(command=self.folder_list.yview)
        
        self.Product_Form()

    def Product_Form(self):

        # Create a style object
        style = ttk.Style()
        
        style.map('BlackOnDisabled.TEntry', foreground=[('disabled', 'black')])

        # Define a custom style named 'Blue.TButton' that changes the foreground color to blue
        style.configure('Blue.TButton', foreground='blue')

        # Create a custom font with a larger size
        link_font = Font(family="Helvetica", size=10)  # Adjust the size as per your requirement
        product_name_font = Font(family="Helvetica", size=11)  # Adjust the size as per your requirement

        # Add validation commands
        validate_percentage_command = (self.register(lambda P: self.validate_input(P, is_percentage=True)), '%P')
        validate_price_command = (self.register(self.validate_input), '%P')
        vcmd = (self.register(self.validate_input), '%P')


        self.product_frame = tk.Frame(self.bottom_frame, bg='light gray')
        self.product_frame.pack(side='right', fill='both', expand=True) #change pack to grid later

        # Row 0 Widgets
        self.row0_frame = tk.Frame(self.product_frame, bg='light gray')
        self.row0_frame.grid(row=0, column=5, sticky='ne', padx=50, pady=0)
    
        self.save_button = ttk.Button(self.row0_frame, text='Save', command=self.save, state='disabled')
        self.save_button.grid(row=0, column=0, sticky='w', padx=0, pady=0)

        self.edit_button = ttk.Button(self.row0_frame, text="Edit", command=self.toggle_edit_mode, state='disabled')
        self.edit_button.grid(row=0, column=1, sticky='w', padx=0, pady=0)


        # Row 1 Widgets
        self.row1_frame = tk.Frame(self.product_frame, bg='light gray')
        self.row1_frame.grid(row=1, column=0, sticky='nw', padx=5, pady=5)
        
        self.order_date_var = tk.StringVar()
        self.order_date_label = ttk.Label(self.row1_frame, text='Order Date')
        self.order_date_label.grid(row=0, column=0, sticky='w', padx=0, pady=0)
        self.order_date_entry = ttk.Entry(self.row1_frame, textvariable=self.order_date_var, state='disabled', style='BlackOnDisabled.TEntry')
        self.order_date_entry.grid(row=1, column=0, sticky='w', padx=0, pady=0)

        self.to_sell_after_var = tk.StringVar()
        self.to_sell_after_label = ttk.Label(self.row1_frame, text='To Sell After')
        self.to_sell_after_label.grid(row=2, column=0, sticky='w', padx=0, pady=0)
        self.to_sell_after_entry = ttk.Entry(self.row1_frame, textvariable=self.to_sell_after_var, state='disabled', style='BlackOnDisabled.TEntry')
        self.to_sell_after_entry.grid(row=3, column=0, sticky='w', padx=0, pady=0)

        # Row 1 Widgets
        # Column 1 Widget
        
        self.r1column1_frame = tk.Frame(self.product_frame, bg='light gray')
        self.r1column1_frame.grid(row=1, column=2, sticky='nw', padx=0, pady=0)
        self.product_image_label = ttk.Label(self.r1column1_frame, text='Image not loaded')
        self.product_image_label.grid(row=0, column=1, sticky='w', padx=0, pady=0)
        
        
        # Row 2 Widgets
        # Column 0 Widget
        # Create a new frame for the column 0 widgets
        self.r2column0_frame = tk.Frame(self.product_frame, bg='light gray')
        self.r2column0_frame.grid(row=2, column=0, sticky='nw', padx=25, pady=25)
        
        self.product_id_var = tk.StringVar()
        self.product_id_label = ttk.Label(self.r2column0_frame, text='Product ID')
        self.product_id_label.grid(row=0, column=0, sticky='w', padx=0, pady=0)
        self.product_id_entry = ttk.Entry(self.r2column0_frame, textvariable=self.product_id_var, state='disabled', style='BlackOnDisabled.TEntry')
        self.product_id_entry.grid(row=1, column=0, sticky='w', padx=0, pady=0)

        self.r2column0_frame.grid_rowconfigure(2, minsize=2)  # Adjust 'minsize' for desired space

        self.product_name_var = tk.StringVar()
        self.product_name_label = ttk.Label(self.r2column0_frame, text='Product Name')
        self.product_name_label.grid(row=3, column=0, sticky='w', padx=0, pady=0)

        # Create the Text widget with the desired background color inside the border frame
        self.product_name_text = tk.Text(self.r2column0_frame, height=8, width=50, bg="#eff0f1", fg="#000000", wrap="word", bd=0, highlightthickness=1, highlightcolor="#94cfeb", font=product_name_font)
        self.product_name_text.grid(row=4, column=0, sticky='w', padx=0, pady=1)
        
        # Bind the mouse click event to an empty lambda function
        self.product_name_text.bind("<Button-1>", lambda e: "break")
        
        self.r2column0_frame.grid_rowconfigure(5, minsize=2)  # Adjust 'minsize' for desired space
        
        self.product_folder_var = tk.StringVar()
        self.product_folder_label = ttk.Label(self.r2column0_frame, text='Product Folder')
        self.product_folder_label.grid(row=6, column=0, sticky='w', padx=0, pady=2)

        # Now use this style when creating your button
        self.product_folder_link = ttk.Button(self.r2column0_frame, textvariable=self.product_folder_var, style='Blue.TButton')

        self.product_folder_link.grid(row=7, column=0, sticky='w', padx=0, pady=0)

        self.r2column0_frame.grid_rowconfigure(8, minsize=2)  # Adjust 'minsize' for desired space

        self.order_link_var = tk.StringVar()
        self.order_link_label = ttk.Label(self.r2column0_frame, text='Order Link')
        self.order_link_label.grid(row=9, column=0, sticky='w', padx=0, pady=0)
        
        # Replace the Entry with a Text widget for clickable links
        self.order_link_text = tk.Text(self.r2column0_frame, height=1, width=40, bg="#eff0f1", fg="#000000", wrap=tk.NONE, bd=0, font=link_font)
        self.order_link_text.grid(row=10, column=0, sticky='w', padx=0, pady=1)
        self.order_link_text.tag_configure("hyperlink", foreground="blue", underline=True)
        self.order_link_text.bind("<Button-1>", self.open_hyperlink)
        self.order_link_text.config(state='disabled')

        self.r2column0_frame.grid_rowconfigure(11, minsize=2)  # Adjust 'minsize' for desired space

        self.asin_var = tk.StringVar()
        self.asin_label = ttk.Label(self.r2column0_frame, text='ASIN')
        self.asin_label.grid(row=12, column=0, sticky='w', padx=0, pady=0)
        self.asin_entry = ttk.Entry(self.r2column0_frame, textvariable=self.asin_var, state='disabled', style='BlackOnDisabled.TEntry')
        self.asin_entry.grid(row=13, column=0, sticky='w', padx=0, pady=0)


        # Row 2 Widgets
        # Column 1 Widgets

        self.r2column1_frame = tk.Frame(self.product_frame, bg='light gray')
        self.r2column1_frame.grid(row=2, column=1, sticky='nw', padx=0, pady=5)
        custom_font = Font(family="Helvetica", size=7)
        style.configure('SmallFont.TButton', font=custom_font, padding=1)
        
        self.r2column1_frame.grid_rowconfigure(0, minsize=75)  # Adjust 'minsize' for desired space
        self.fair_market_value_var = tk.StringVar()
        self.fair_market_value_label = ttk.Label(self.r2column1_frame, text='Fair Market Value')
        self.fair_market_value_label.grid(row=2, column=0, sticky='w', padx=0, pady=0)
        self.fair_market_value_entry = ttk.Entry(self.r2column1_frame, textvariable=self.fair_market_value_var, state='disabled', style='BlackOnDisabled.TEntry')
        self.fair_market_value_entry.grid(row=3, column=0, sticky='w', padx=0, pady=0)
        
        self.regular_product_price_var = tk.StringVar()
        self.regular_product_price_label = ttk.Label(self.r2column1_frame, text='Product Price')
        self.regular_product_price_label.grid(row=4, column=0, sticky='w', padx=0, pady=0)
        self.regular_product_price_entry = ttk.Entry(self.r2column1_frame, textvariable=self.regular_product_price_var, state='disabled', style='BlackOnDisabled.TEntry')
        self.regular_product_price_entry.grid(row=5, column=0, sticky='w', padx=0, pady=0)
        
        self.ivu_tax_var = tk.StringVar()
        self.ivu_tax_label = ttk.Label(self.r2column1_frame, text='IVU Tax')
        self.ivu_tax_label.grid(row=6, column=0, sticky='w', padx=0, pady=0)
        self.ivu_tax_entry = ttk.Entry(self.r2column1_frame, textvariable=self.ivu_tax_var, state='disabled', style='BlackOnDisabled.TEntry')
        self.ivu_tax_entry.grid(row=7, column=0, sticky='w', padx=0, pady=0)
        
        self.product_price_plus_ivu_var = tk.StringVar()
        self.product_price_plus_ivu_label = ttk.Label(self.r2column1_frame, text='Product Price (+ IVU)')
        self.product_price_plus_ivu_label.grid(row=8, column=0, sticky='w', padx=0, pady=0)
        self.product_price_plus_ivu_entry = ttk.Entry(self.r2column1_frame, textvariable=self.product_price_plus_ivu_var, state='disabled', style='BlackOnDisabled.TEntry')
        self.product_price_plus_ivu_entry.grid(row=9, column=0, sticky='w', padx=0, pady=0)

        # Row 2 Widgets
        # Column 2 Widgets

        self.r2column2_frame = tk.Frame(self.product_frame, bg='light gray')
        self.r2column2_frame.grid(row=2, column=2, sticky='nw', padx=0, pady=5)
        custom_font = Font(family="Helvetica", size=7)
        style.configure('SmallFont.TButton', font=custom_font, padding=1)
        
        self.r2column2_frame.grid_rowconfigure(0, minsize=75)  # Adjust 'minsize' for desired space

        self.discount_var = tk.StringVar()
        self.discount_label = ttk.Label(self.r2column2_frame, text='Discount($ Or %)')
        self.discount_label.grid(row=1, column=0, sticky='w', padx=0, pady=0)

        # Frame to hold the discount entries
        self.discount_frame = ttk.Frame(self.r2column2_frame)
        self.discount_frame.grid(row=2, column=0, sticky='w', padx=0, pady=0)

        # Discount entries with validation and event binding
        self.discount_var = tk.StringVar()
        self.discount_entry = ttk.Entry(self.discount_frame, textvariable=self.discount_var, width=8, state='disabled', style='BlackOnDisabled.TEntry', validate='key', validatecommand=validate_price_command)
        self.discount_entry.pack(side=tk.LEFT)
        #self.discount_entry.bind("<KeyRelease>", self.on_price_changed)        
        self.discount_entry.bind("<FocusIn>", self.on_discount_price_focus_in)        
        self.discount_entry.bind("<FocusOut>", self.on_discount_price_focus_out)

        # Label "Or"
        self.or_label = ttk.Label(self.discount_frame, text="Or")
        self.or_label.pack(side=tk.LEFT)

        self.percent_discount_var = tk.StringVar()
        self.percent_discount_entry = ttk.Entry(self.discount_frame, textvariable=self.percent_discount_var, width=8, state='disabled', style='BlackOnDisabled.TEntry', validate='key', validatecommand=validate_percentage_command)
        self.percent_discount_entry.pack(side=tk.LEFT)
        #self.percent_discount_entry.bind("<KeyRelease>", self.on_percentage_changed)
        self.percent_discount_entry.bind("<FocusIn>", self.on_discount_percentage_focus_in)
        self.percent_discount_entry.bind("<FocusOut>", self.on_discount_percentage_focus_out)
        
        self.product_price_after_discount_var = tk.StringVar()
        self.product_price_after_discount_label = ttk.Label(self.r2column2_frame, text='Product Price after Discount')
        self.product_price_after_discount_label.grid(row=3, column=0, sticky='w', padx=0, pady=0)
        self.product_price_after_discount_entry = ttk.Entry(self.r2column2_frame, textvariable=self.product_price_after_discount_var, state='disabled', style='BlackOnDisabled.TEntry')
        self.product_price_after_discount_entry.grid(row=4, column=0, sticky='w', padx=0, pady=0)

        self.ivu_tax_after_discount_var = tk.StringVar()
        self.ivu_tax_after_discount_label = ttk.Label(self.r2column2_frame, text='IVU Tax after Discount')
        self.ivu_tax_after_discount_label.grid(row=5, column=0, sticky='w', padx=0, pady=0)
        self.ivu_tax_after_discount_entry = ttk.Entry(self.r2column2_frame, textvariable=self.ivu_tax_after_discount_var, state='disabled', style='BlackOnDisabled.TEntry')
        self.ivu_tax_after_discount_entry.grid(row=6, column=0, sticky='w', padx=0, pady=0)

        self.product_price_minus_discount_plus_ivu_var = tk.StringVar()
        self.product_price_minus_discount_plus_ivu_label = ttk.Label(self.r2column2_frame, text='Product Price (+IVU - Discount)')
        self.product_price_minus_discount_plus_ivu_label.grid(row=7, column=0, sticky='w', padx=0, pady=0)
        self.product_price_minus_discount_plus_ivu_entry = ttk.Entry(self.r2column2_frame, textvariable=self.product_price_minus_discount_plus_ivu_var, state='disabled', style='BlackOnDisabled.TEntry')
        self.product_price_minus_discount_plus_ivu_entry.grid(row=8, column=0, sticky='w', padx=0, pady=0)

        self.sold_date_var = tk.StringVar()
        self.sold_date_label = ttk.Label(self.r2column2_frame, text='Sold Date')
        self.sold_date_label.grid(row=9, column=0, sticky='w', padx=0, pady=0)
        
        self.sold_date_entry = ttk.Entry(self.r2column2_frame, textvariable=self.sold_date_var, state='disabled', style='BlackOnDisabled.TEntry')
        self.sold_date_entry.grid(row=10, column=0, sticky='w', padx=0, pady=0)
        
        self.sold_date_button = ttk.Button(self.r2column2_frame, text="Pick\nDate", style='SmallFont.TButton', command=self.pick_date, state='disabled', width=5)
        self.sold_date_button.grid(row=10, column=0, sticky='e', padx=0, pady=0)

        self.clear_button = ttk.Button(self.r2column2_frame, text="Clear\nDate", style='SmallFont.TButton', command=self.clear_date, state='disabled', width=5)
        self.clear_button.grid(row=10, column=1, sticky='e', padx=0, pady=0)

        self.payment_type_var = tk.StringVar()
        self.payment_type_label = ttk.Label(self.r2column2_frame, text='Payment Type')
        self.payment_type_label.grid(row=11, column=0, sticky='w', padx=0, pady=0)
        
        self.payment_type_combobox = ttk.Combobox(self.r2column2_frame, textvariable=self.payment_type_var, state='disabled', style='BlackOnDisabled.TEntry')
        self.payment_type_combobox['values'] = ('', 'Cash', 'ATH Movil')
        self.payment_type_combobox.grid(row=12, column=0, sticky='w', padx=0, pady=0)        
        
        self.sold_price_var = tk.StringVar()
        self.sold_price_label = ttk.Label(self.r2column2_frame, text='Sold Price')
        self.sold_price_label.grid(row=13, column=0, sticky='w', padx=0, pady=0)
        self.sold_price_entry = ttk.Entry(self.r2column2_frame, textvariable=self.sold_price_var, state='disabled', style='BlackOnDisabled.TEntry')
        self.sold_price_entry.grid(row=14, column=0, sticky='w', padx=0, pady=0)


        # Row 2 Widgets
        # Column 3 Widgets
        # Creating a new frame for checkboxes within the product frame
        self.checkbox_frame = tk.Frame(self.product_frame, bg='light gray')
        self.checkbox_frame.grid(row=2, column=3, rowspan=8, sticky='nw', padx=0, pady=5)
        self.checkbox_frame.grid_rowconfigure(0, minsize=75)  # Adjust 'minsize' for desired space

        self.sold_var = tk.BooleanVar()
        self.sold_checkbutton = ttk.Checkbutton(self.checkbox_frame, text='Sold', variable=self.sold_var)
        self.sold_checkbutton.grid(row=1, column=0, sticky='w', padx=0, pady=0)
        
        self.checkbox_frame.grid_rowconfigure(2, minsize=20)  # This creates a 20-pixel-high empty row as a spacer
        
        self.cancelled_order_var = tk.BooleanVar()
        self.cancelled_order_checkbutton = ttk.Checkbutton(self.checkbox_frame, text='Cancelled Order', variable=self.cancelled_order_var)
        self.cancelled_order_checkbutton.grid(row=3, column=0, sticky='w', padx=0, pady=0)

        self.damaged_var = tk.BooleanVar()
        self.damaged_checkbutton = ttk.Checkbutton(self.checkbox_frame, text='Damaged', variable=self.damaged_var)
        self.damaged_checkbutton.grid(row=4, column=0, sticky='w', padx=0, pady=0)

        self.personal_var = tk.BooleanVar()
        self.personal_checkbutton = ttk.Checkbutton(self.checkbox_frame, text='Personal', variable=self.personal_var)
        self.personal_checkbutton.grid(row=5, column=0, sticky='w', padx=0, pady=0)

        self.checkbox_frame.grid_rowconfigure(6, minsize=20)  # This creates a 20-pixel-high empty row as a spacer

        self.reviewed_var = tk.BooleanVar()
        self.reviewed_checkbutton = ttk.Checkbutton(self.checkbox_frame, text='Reviewed', variable=self.reviewed_var)
        self.reviewed_checkbutton.grid(row=7, column=0, sticky='w', padx=0, pady=0)

        self.pictures_downloaded_var = tk.BooleanVar()
        self.pictures_downloaded_checkbutton = ttk.Checkbutton(self.checkbox_frame, text='Pictures Downloaded', variable=self.pictures_downloaded_var)
        self.pictures_downloaded_checkbutton.grid(row=8, column=0, sticky='w', padx=0, pady=0)


        self.product_frame.grid_rowconfigure(3, minsize=60)  # This creates a 20-pixel-high empty row as a spacer
        

        # Row 4 Widgets
        # Column 0 Widgets
        # Creating a new frame for checkboxes within the product frame
        self.comments_frame = tk.Frame(self.product_frame, bg='light gray')
        self.comments_frame.grid(row=4, column=0, columnspan=3, sticky='nw', padx=25, pady=5)

        self.comments_text = tk.Text(self.comments_frame, height=8, width=150, bg="#eff0f1", fg="#000000", wrap="word", state="disabled", bd=0, highlightthickness=1, highlightcolor="#94cfeb", font=product_name_font)
        self.comments_text.grid(row=4, column=0, sticky='w', padx=0, pady=1)


        # Bind the new checkbox click control function to the checkboxes
        self.sold_checkbutton.bind('<Button-1>', lambda e: self.checkbox_click_control(self.sold_var))
        self.cancelled_order_checkbutton.bind('<Button-1>', lambda e: self.checkbox_click_control(self.cancelled_order_var))
        self.damaged_checkbutton.bind('<Button-1>', lambda e: self.checkbox_click_control(self.damaged_var))
        self.personal_checkbutton.bind('<Button-1>', lambda e: self.checkbox_click_control(self.personal_var))
        self.reviewed_checkbutton.bind('<Button-1>', lambda e: self.checkbox_click_control(self.reviewed_var))
        self.pictures_downloaded_checkbutton.bind('<Button-1>', lambda e: self.checkbox_click_control(self.pictures_downloaded_var))

        # Add focus in and focus out bindings for price-related entry fields
        self.fair_market_value_entry.bind("<FocusIn>", self.on_price_focus_in)
        self.fair_market_value_entry.bind("<FocusOut>", self.on_price_focus_out)

        self.regular_product_price_entry.bind("<FocusIn>", self.on_price_focus_in)
        self.regular_product_price_entry.bind("<FocusOut>", self.on_price_focus_out)

        self.ivu_tax_entry.bind("<FocusIn>", self.on_price_focus_in)
        self.ivu_tax_entry.bind("<FocusOut>", self.on_price_focus_out)

        self.product_price_plus_ivu_entry.bind("<FocusIn>", self.on_price_focus_in)
        self.product_price_plus_ivu_entry.bind("<FocusOut>", self.on_price_focus_out)

        self.product_price_after_discount_entry.bind("<FocusIn>", self.on_price_focus_in)
        self.product_price_after_discount_entry.bind("<FocusOut>", self.on_price_focus_out)

        self.ivu_tax_after_discount_entry.bind("<FocusIn>", self.on_price_focus_in)
        self.ivu_tax_after_discount_entry.bind("<FocusOut>", self.on_price_focus_out)

        self.product_price_minus_discount_plus_ivu_entry.bind("<FocusIn>", self.on_price_focus_in)
        self.product_price_minus_discount_plus_ivu_entry.bind("<FocusOut>", self.on_price_focus_out)

        self.sold_price_entry.bind("<FocusIn>", self.on_price_focus_in)
        self.sold_price_entry.bind("<FocusOut>", self.on_price_focus_out)

        # configure validation commands
        self.fair_market_value_entry.config(validate='key', validatecommand=vcmd)
        self.regular_product_price_entry.config(validate='key', validatecommand=vcmd)
        self.ivu_tax_entry.config(validate='key', validatecommand=validate_price_command)
        self.product_price_plus_ivu_entry.config(validate='key', validatecommand=vcmd)
        self.product_price_after_discount_entry.config(validate='key', validatecommand=vcmd)
        self.ivu_tax_after_discount_entry.config(validate='key', validatecommand=vcmd)
        self.product_price_minus_discount_plus_ivu_entry.config(validate='key', validatecommand=vcmd)
        self.sold_price_entry.config(validate='key', validatecommand=vcmd)

        # Load settings
        try:
            with open("folders_paths.txt", "r") as file:
                lines = file.read().splitlines()
                self.inventory_folder = lines[0]
                self.sold_folder = lines[1]
                self.to_sell_folder = lines[2] if len(lines) > 2 else None
                if self.inventory_folder:  # Check if inventory_folder is defined
                    self.combine_and_display_folders()
        except FileNotFoundError:
            pass
        self.search_entry.focus_set()

    def validate_input(self, input_value, is_percentage=False):
        """Validates the input to allow only one decimal point and up to two decimal places."""
        if input_value == "":
            return True

        if input_value.count('.') > 1:
            return False

        parts = input_value.split('.')
        if len(parts) == 2 and len(parts[1]) > 2:
            return False

        return all(ch.isdigit() or ch == '.' for ch in input_value)

    def on_price_focus_in(self, event):
        """Stores the initial price when focus is gained."""
        entry_widget = event.widget
        price_str = entry_widget.get()
        if event.widget == self.product_price_plus_ivu_entry:
            # Only store the price for the specific field
            self.initial_product_price_plus_ivu = price_str.lstrip('$')
        if price_str.startswith('$'):
            entry_widget.delete(0, tk.END)
            entry_widget.insert(0, price_str.lstrip('$'))

    def on_price_focus_out(self, event):
        if self.edit_mode:
            if self.trigger_price_focus_out_flag:

                entry_widget = event.widget
                current_price = entry_widget.get().lstrip('$')

                entry_widget.config(validate='none')

                # Check if the price string is empty and set it and related fields to default values
                try:
                    if not current_price.strip():
                        # Temporarily disable validation and update widget state
                        entry_list = [self.regular_product_price_entry, self.ivu_tax_entry, 
                                    self.discount_entry, self.product_price_after_discount_entry, 
                                    self.ivu_tax_after_discount_entry, 
                                    self.product_price_minus_discount_plus_ivu_entry, 
                                    self.percent_discount_entry]

                        for widget in entry_list:
                            widget.config(validate='none', state='normal')

                            # Set the current widget to $0
                            self.product_price_plus_ivu_entry.delete(0, tk.END)
                            self.product_price_plus_ivu_entry.insert(0, "$0")

                            # Set related fields to their default values
                            
                            self.regular_product_price_entry.delete(0, tk.END)
                            self.regular_product_price_entry.insert(0, "$0")

                            self.ivu_tax_entry.delete(0, tk.END)
                            self.ivu_tax_entry.insert(0, "$0")

                            # Temporarily disable validation and change state to normal
                            self.discount_entry.config(validate='none', state='normal')

                            # Set the widget to '$0'
                            self.discount_entry.delete(0, tk.END)
                            self.discount_entry.insert(0, "$0")

                            # Re-enable validation and change state back to disabled
                            self.discount_entry.config(validate='key', state='disabled')


                            self.product_price_after_discount_entry.delete(0, tk.END)
                            self.product_price_after_discount_entry.insert(0, "$0")

                            self.ivu_tax_after_discount_entry.delete(0, tk.END)
                            self.ivu_tax_after_discount_entry.insert(0, "$0")

                            self.product_price_minus_discount_plus_ivu_entry.delete(0, tk.END)
                            self.product_price_minus_discount_plus_ivu_entry.insert(0, "$0")

                            self.percent_discount_entry.delete(0, tk.END)
                            self.percent_discount_entry.insert(0, "0%")

                            # Re-enable validation and update widget state, disable all except specific widgets
                            for widget in entry_list:
                                if widget not in [self.discount_entry, self.percent_discount_entry, self.product_price_plus_ivu_entry]:
                                    widget.config(validate='key', state='disabled')
                                else:
                                    widget.config(validate='key', state='normal')  # Keep these widgets editable

                        # Update GUI
                        self.update_idletasks()

                        return           
                except Exception as e:
                    print(f"Error while setting prices as $0/0%: {e}")
                
                if not current_price.startswith('$'):
                    entry_widget.delete(0, tk.END)
                    entry_widget.insert(0, f"${current_price}")

                entry_widget.config(validate='key')

                if event.widget == self.product_price_plus_ivu_entry and self.initial_product_price_plus_ivu != current_price:
                    if not hasattr(self, 'prompt_shown'):
                        self.prompt_shown = True

                        self.recalculate_original_price_and_tax()
                        discount_price = self.discount_var.get().lstrip('$')
                        discount_percentage = self.percent_discount_var.get().rstrip('%')
                        message = f"Product price changed. Calculate discount based on?\n\nPrice: ${discount_price}\nPercentage: {discount_percentage}%"
                        response = messagebox.askquestion("Discount Calculation", message)

                        if response == 'yes':
                            self.calculate_discount('price')
                            # if self.trigger_save_flag:
                            #     self.save()
                            #     self.trigger_save_flag = False

                        else:
                            
                            self.calculate_discount('percentage')
                            # if self.trigger_save_flag:
                            #     self.save()
                            #     self.trigger_save_flag = False

                        del self.prompt_shown

                    self.initial_product_price_plus_ivu = ''

    def save_on_key_handler(self, event):
    
        if event.widget in [self.discount_entry, self.percent_discount_entry, self.product_price_plus_ivu_entry]:
            #self.trigger_save_flag = True 
            self.product_id_entry.focus_set()
        else:
            self.save()

    def edit_on_key_handler(self, event):

        productid = self.product_id_entry.get().upper()
        if event.widget in [self.discount_entry, self.percent_discount_entry, self.product_price_plus_ivu_entry]:
            self.trigger_price_focus_out_flag = False
            self.search_entry.focus_set()
            self.trigger_price_focus_out_flag = True
            self.refresh_and_select_product(productid)
            self.toggle_edit_mode()
        else:
            self.toggle_edit_mode()

    def on_price_changed(self, event=None):
        self.last_changed = 'price'
        self.calculate_discount()

    def on_discount_price_focus_in(self, event=None):
        """Removes '$' symbol and stores the rounded value of the discount price when focus is gained."""
        price_str = self.discount_var.get()
        if price_str.startswith('$'):
            price_str = price_str.lstrip('$')
            self.discount_var.set(price_str)
        
        # Store the rounded numerical value
        try:
            self.initial_discount_price = round(float(price_str), 2)
        except ValueError:
            self.initial_discount_price = None

    def on_discount_price_focus_out(self, event=None):
        """Adds '$' symbol to the discount price when focus is lost."""
        price_str = self.discount_var.get()

        # Check if the price string is empty or invalid, set it to '$0'
        if not price_str or not price_str.replace('$', '').strip().replace('.', '', 1).isdigit():
            self.discount_var.set("$0")
            final_discount_price = 0.0
        else:
            if not price_str.startswith('$'):
                self.discount_var.set(f"${price_str}")

            try:
                final_discount_price = round(float(price_str.lstrip('$')), 2)
            except ValueError:
                final_discount_price = None

        # Trigger discount calculation only if the price has changed
        if self.initial_discount_price != final_discount_price:
            self.last_changed = 'price'
            self.calculate_discount('price')  # Pass 'price' as the argument
        else:
            # Optionally, handle the case where the price hasn't changed
            pass

    def on_percentage_changed(self, *args):
        self.last_changed = 'percentage'
        self.calculate_discount()

    def on_discount_percentage_focus_in(self, event=None):
        """Removes '%' symbol from the discount percentage when focus is gained."""
        percentage_str = self.percent_discount_var.get()

        if percentage_str.endswith('%'):
            percentage_str = percentage_str.rstrip('%')
            self.percent_discount_var.set(percentage_str)
        
        # Now try converting the stripped string to a float
        try:
            self.initial_percent_discount = round(float(percentage_str), 2)
        except ValueError:
            self.initial_percent_discount = None

    def on_discount_percentage_focus_out(self, event=None):
        """Adds '%' symbol to the discount percentage when focus is lost."""
        percentage_str = self.percent_discount_var.get()

        # Check if the percentage string is empty or invalid, set it to '0%'
        if not percentage_str or not percentage_str.replace('%', '').strip().isdigit():
            self.percent_discount_var.set("0%")
            final_percent_discount = 0.0
        else:
            if not percentage_str.endswith('%'):
                self.percent_discount_var.set(f"{percentage_str}%")

            try:
                final_percent_discount = round(float(percentage_str.strip('%')), 2)
            except ValueError:
                final_percent_discount = None

        # Trigger discount calculation only if the percentage has changed
        if self.initial_percent_discount != final_percent_discount:
            self.last_changed = 'percentage'
            self.calculate_discount('percentage')  # Pass 'percentage' as the argument
        else:
            # Optionally, handle the case where the percentage hasn't changed
            pass


        #if #flag
            #self.save()
            #flag = false

    def custom_float_format(self, value):
        """Formats the float value to string with two decimal places."""
        return "{:.2f}".format(value)

    def calculate_discount(self, based_on):
        try:
            price_str = self.regular_product_price_var.get().lstrip('$')
            price = Decimal(price_str) if price_str else Decimal('0')

            if based_on == 'percentage':
                percentage_str = self.percent_discount_var.get().strip('%')
                percentage = Decimal(percentage_str) if percentage_str else Decimal('0')
                discount_price = (price * percentage / Decimal('100')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
                self.discount_var.set(f"${discount_price:.2f}")

            elif based_on == 'price':
                discount_str = self.discount_var.get().strip('$')
                discount = Decimal(discount_str) if discount_str else Decimal('0')
                percentage = ((discount / price) * Decimal('100')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
                
                # Adjust the format to allow for decimal percentages
                formatted_percentage = "{:.2f}%".format(percentage)
                self.percent_discount_var.set(formatted_percentage)

        except (ValueError, InvalidOperation):
            pass

        self.calculate_discount_fields()

    def calculate_discount_fields(self):
        # Helper function to strip characters and convert to Decimal
        def clean_and_convert(value, strip_char=None):
            if strip_char:
                value = value.replace(strip_char, '')
            try:
                return Decimal(value)
            except (ValueError, InvalidOperation):
                return Decimal('0')

        # Get values and clean them
        product_price_plus_ivu = clean_and_convert(self.product_price_plus_ivu_var.get(), '$')
        discount_price = clean_and_convert(self.discount_var.get(), '$')
        discount_percentage = clean_and_convert(self.percent_discount_var.get(), '%')

        # Define the tax rate
        tax_rate = Decimal('0.115')

        # Correctly calculate the original product price and IVU Tax
        original_product_price = (product_price_plus_ivu / (1 + tax_rate)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        ivu_tax = (original_product_price * tax_rate).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

        # Check if there's a discount
        has_discount = discount_price > 0 or discount_percentage > 0

        # Calculate Discounted Prices
        if has_discount:
            # Determine the discount amount
            if discount_price > 0:
                discount_amount = discount_price
            else:
                discount_amount = (original_product_price * (discount_percentage / Decimal('100'))).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

            # Apply the discount to the original product price
            product_price_after_discount = (original_product_price - discount_amount).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

            # Recalculate the IVU tax based on the discounted price
            ivu_tax_after_discount = (product_price_after_discount * tax_rate).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

            # Calculate the total price after discount including IVU tax
            product_price_plus_ivu_discount = (product_price_after_discount + ivu_tax_after_discount).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        else:
            product_price_after_discount = original_product_price
            ivu_tax_after_discount = ivu_tax
            product_price_plus_ivu_discount = product_price_plus_ivu

        # Update the fields
        self.product_price_after_discount_var.set(f"${product_price_after_discount:.2f}")
        self.ivu_tax_after_discount_var.set(f"${ivu_tax_after_discount:.2f}")
        self.product_price_minus_discount_plus_ivu_var.set(f"${product_price_plus_ivu_discount:.2f}")

    def recalculate_original_price_and_tax(self):
        # Extract and clean the product price (+ IVU) value
        price_plus_ivu_str = self.product_price_plus_ivu_var.get().lstrip('$')
        try:
            price_plus_ivu = Decimal(price_plus_ivu_str)
        except ValueError:
            price_plus_ivu = Decimal('0')

        # Define the tax rate
        tax_rate = Decimal('0.115')

        # Calculate the original product price before tax
        original_product_price = (price_plus_ivu / (1 + tax_rate)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

        # Calculate the IVU tax based on the original product price
        IVU_tax = (original_product_price * tax_rate).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)


        price_plus_ivu_str = original_product_price + IVU_tax
        # Update the IVU tax and original product price fields
        self.ivu_tax_var.set(f"${IVU_tax:.2f}")
        self.regular_product_price_var.set(f"${original_product_price:.2f}")
        self.product_price_plus_ivu_var.set(f"${price_plus_ivu_str:.2f}")

    def pick_date(self):
        def grab_date():
            selected_date = cal.selection_get()  # Get the selected date
            formatted_date = selected_date.strftime('%m/%d/%Y')  # Format the date

            self.sold_date_entry.config(state="normal")  # Enable the entry widget
            self.sold_date_entry.delete(0, tk.END)  # Clear the entry field
            self.sold_date_entry.insert(0, formatted_date)  # Insert the formatted date
            self.sold_date_entry.config(state="disabled")  # Disable the entry widget

            top.destroy()  # Close the Toplevel window
        def select_today_and_close(event):
            cal.selection_set(datetime.today())  # Set selection to today's date
            grab_date()  # Then grab the date and close
        top = tk.Toplevel(self)
        today = datetime.today()
        cal = Calendar(top, selectmode='day', year=today.year, month=today.month, day=today.day)
        cal.pack(pady=20)    # Set focus to the Toplevel window and bind the Enter key
        top.focus_set()
        top.bind('<Return>', select_today_and_close)
        cal.bind("<<CalendarSelected>>", lambda event: grab_date())
    
    def clear_date(self):
        self.sold_date_entry.config(state="normal")  # Enable the entry widget
        self.sold_date_entry.delete(0, tk.END)  # Clear the entry field
        self.sold_date_entry.config(state="disabled")  # Disable the entry widget

    def open_hyperlink(self, event):
        try:
            start_index = self.order_link_text.index("@%s,%s" % (event.x, event.y))
            tag_indices = list(self.order_link_text.tag_ranges('hyperlink'))
            for start, end in zip(tag_indices[0::2], tag_indices[1::2]):
                if self.order_link_text.compare(start_index, ">=", start) and self.order_link_text.compare(start_index, "<=", end):
                    url = self.order_link_text.get(start, end)
                    webbrowser.open(url)
                    return "break"
        except Exception as e:
            print(f"Error when opening hyperlink: {e}")

    def Settings_Window_Start(self):
        if hasattr(self, 'settings_window') and self.settings_window.winfo_exists():
            self.settings_window.lift()
            return
        self.settings_window = tk.Toplevel(self)
        self.settings_window.title("Settings")
        self.settings_window.state('zoomed')

        # Create and grid the settings frame
        self.settings_frame = tk.Frame(self.settings_window)
        self.settings_frame.grid(row=1, column=1, sticky='nw')

        # Load settings
        self.default_filepath, self.default_sheet = self.load_excel_settings()
        
        # Configure the grid columns of the frame
        self.settings_frame.grid_columnconfigure(1, weight=1)  # Adjust the weight as needed

        # Now grid all widgets onto the settings_frame
        self.inventory_folder_button = ttk.Button(self.settings_frame, text="Choose Inventory Folder", command=self.choose_inventory_folder)
        self.inventory_folder_button.grid(row=1, column=0, padx=5, pady=5, sticky='w')
        self.inventory_folder_label = ttk.Label(self.settings_frame, text=self.inventory_folder if self.inventory_folder else "Not chosen")
        self.inventory_folder_label.grid(row=1, column=1, padx=5, pady=5, sticky='w')

        self.sold_folder_button = ttk.Button(self.settings_frame, text="Choose Sold Inventory Folder", command=self.choose_sold_folder)
        self.sold_folder_button.grid(row=2, column=0, padx=5, pady=5, sticky='w')
        self.sold_folder_label = ttk.Label(self.settings_frame, text=self.sold_folder if self.sold_folder else "Not chosen")
        self.sold_folder_label.grid(row=2, column=1, padx=5, pady=5, sticky='w')

        self.to_sell_folder_button = ttk.Button(self.settings_frame, text="Choose Products to Sell Folder", command=self.choose_to_sell_folder)
        self.to_sell_folder_button.grid(row=3, column=0, padx=5, pady=5, sticky='w')
        self.to_sell_folder_label = ttk.Label(self.settings_frame, text=self.to_sell_folder if self.to_sell_folder else "Not chosen")
        self.to_sell_folder_label.grid(row=3, column=1, padx=5, pady=5, sticky='w')

        self.excel_db_button = ttk.Button(self.settings_frame, text="Select Excel Database", command=self.select_excel_database)
        self.excel_db_button.grid(row=4, column=0, padx=5, pady=5, sticky='w')
        excel_db_text = f"{self.default_filepath} - Sheet: {self.default_sheet}" if self.default_filepath and self.default_sheet else "Not chosen"
        self.excel_db_label = ttk.Label(self.settings_frame, text=excel_db_text)
        self.excel_db_label.grid(row=4, column=1, padx=5, pady=5, sticky='w')

        self.create_word_files_button = ttk.Button(self.settings_frame, text="Create Word Files for Products", command=self.correlate_data)
        self.create_word_files_button.grid(row=5, column=0, padx=5, pady=5, sticky='w')

        self.autofill_links_asin_tosellafter_data_button = ttk.Button(self.settings_frame, text="Autofill Excel Data(link, asin, tosellafter)", command=self.update_links_in_excel)
        self.autofill_links_asin_tosellafter_data_button.grid(row=6, column=0, padx=5, pady=5, sticky='w')

        self.update_foldersnames_folderpaths_button = ttk.Button(self.settings_frame, text="Update folder names and paths", command=self.update_folders_paths)
        self.update_foldersnames_folderpaths_button.grid(row=7, column=0, padx=5, pady=5, sticky='w')

        self.products_to_sell_list_button = ttk.Button(self.settings_frame, text="Show list of products available to sell", command=self.products_to_sell_report)
        self.products_to_sell_list_button.grid(row=8, column=0, padx=5, pady=5, sticky='w')

        self.update_prices_button = ttk.Button(self.settings_frame, text="Update empty product prices based on Fair Market Value.", command=self.update_prices)
        self.update_prices_button.grid(row=9, column=0, padx=5, pady=5, sticky='w')

        self.back_button = ttk.Button(self.settings_window, text="<- Back", command=self.back_to_main)
        self.back_button.grid(row=0, column=0, sticky='w', padx=5, pady=5)

        self.combine_and_display_folders()
        self.settings_window.protocol("WM_DELETE_WINDOW", lambda: on_close(self, self.master))

        self.master.withdraw()
   
    def on_settings_close(self):
        self.master.destroy()
    
    def get_previous_excel_report_data(self):
        to_sell_folder = self.to_sell_folder
        latest_file_date = None
        latest_file = None

        # Today's date for comparison
        today = datetime.today().date()

        # List all files and find the latest one
        for file in os.listdir(to_sell_folder):
            if file.startswith("Products To Sell -") and file.endswith(".xlsx"):
                # Correct the indices to exclude the leading space
                file_date_str = file[len("Products To Sell - "):-len(".xlsx")]
                print(f"Extracted date string: '{file_date_str}' from file name: '{file}'")  # Debugging
                try:
                    file_date = datetime.strptime(file_date_str, "%Y-%m-%d").date()
                    print(f"File: {file}, File Date: {file_date}, Today: {today}")  # Debugging
                    if file_date < today and (latest_file_date is None or file_date > latest_file_date):
                        latest_file_date = file_date
                        latest_file = file
                except ValueError as e:
                    print(f"Error parsing date from file name: {e}")  # Debugging
                    # If the date format is incorrect, skip this file
                    continue

        # Check if a file was found
        if latest_file is None:
            return [0], None  # Return None for latest_file if not found


        # Read the Excel file and get Product IDs
        file_path = os.path.join(to_sell_folder, latest_file)
        workbook = load_workbook(file_path, data_only=True)
        sheet = workbook.active
        product_ids = []
        for row in sheet.iter_rows(min_row=2, values_only=True):  # Assuming Product IDs start from the second row
            product_id = row[0]  # Assuming Product IDs are in the first column
            product_ids.append(product_id)

        return product_ids, latest_file_date  # Return both product_ids and the latest file

    def products_to_sell_report(self):

        # Ensure the Excel file path and sheet name are set
        filepath, sheet_name = self.load_excel_settings()
        if not filepath or not sheet_name:
            messagebox.showerror("Error", "Excel file path or sheet name is not set.")
            return

        # Define the To Sell folder path
        to_sell_folder = self.to_sell_folder
        if not os.path.exists(to_sell_folder):
            messagebox.showerror("Error", "To Sell folder path is not set or does not exist.")
            return

        # Load the original workbook and read the specified sheet into a DataFrame
        workbook = load_workbook(filepath, data_only=True)
        sheet = workbook[sheet_name]
        data = sheet.values
        columns = next(data)[0:]
        df = pd.DataFrame(data, columns=columns)

        # Filter out unwanted products and keep only necessary columns
        df = df[(df['Damaged'] != 'YES') & (df['Cancelled Order'] != 'YES') & (df['Personal'] != 'YES') & (df['Sold'] != 'YES') & (~pd.isna(df['Product ID']))]
        df = df[['Product ID', 'To Sell After', 'Product Name', 'Fair Market Value']]

        # Convert 'To Sell After' to datetime
        df['To Sell After'] = pd.to_datetime(df['To Sell After'], errors='coerce')
        today = pd.to_datetime('today').normalize()
        df = df.dropna(subset=['To Sell After'])
        df = df[df['To Sell After'] <= today]

        # Sort the DataFrame
        sorted_df = df.sort_values(by='To Sell After', ascending=False)

        # Call get_previous_excel_report_data and assign the return value to listx
        previous_product_ids, latest_file_date = self.get_previous_excel_report_data()

        # Define the light green fill
        light_green_fill = PatternFill(start_color='90EE90', end_color='90EE90', fill_type='solid')

        # Create a new workbook and add the sorted data to it
        new_workbook = Workbook()
        new_sheet = new_workbook.active
        new_sheet.title = sheet_name

        for r_idx, row in enumerate(dataframe_to_rows(sorted_df, index=False, header=True), start=1):
            for c_idx, value in enumerate(row, start=1):
                cell = new_sheet.cell(row=r_idx, column=c_idx, value=value)

                if c_idx == 1 and r_idx > 1:  # Skip header row                    
                    if cell.value is not None and cell.value.upper() in previous_product_ids:
                        previous_product_ids.remove(cell.value)
                    else:
                        cell.fill = light_green_fill
                if c_idx == 2 and r_idx > 1:  # Skip header row
                    cell.number_format = 'MM/DD/YYYY'
                # Apply currency format to 'Fair Market Value' column (assuming it's the fourth column)
                if c_idx == 4 and r_idx > 1:  # Skip header row
                    cell.number_format = '"$"#,##0.00'
                # Apply middle and center alignment to all cells
                if c_idx == 3 and r_idx > 1:  # 'Product Name' column
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                else:
                    cell.alignment = Alignment(horizontal='center', vertical='center')

        # Define the table dimensions
        table_ref = f"A1:{chr(65 + sorted_df.shape[1] - 1)}{sorted_df.shape[0] + 1}"

        # Create a table
        table = Table(displayName="ProductsToSellTable", ref=table_ref)
        style = TableStyleInfo(name="TableStyleMedium9", showFirstColumn=False,
                            showLastColumn=False, showRowStripes=True, showColumnStripes=True)
        table.tableStyleInfo = style
        new_sheet.add_table(table)

        # Adjust column widths
        new_sheet.column_dimensions['A'].width = 120 / 7  # Width for 'Product ID'
        new_sheet.column_dimensions['B'].width = 120 / 7  # Width for 'To Sell After'
        new_sheet.column_dimensions['C'].width = 700 / 7  # Width for 'Product Name'
        new_sheet.column_dimensions['D'].width = 120 / 7  # Width for 'Fair Market Value'

        # Assigning values to cells# Assuming latest_file is a string in the format "YYYY-MM-DD"
        formatted_date = latest_file_date.strftime('%A, %B %d, %Y')
        new_sheet['F2'] = f"Product IDs highlighted in green represent new products added since the \nlast report from {formatted_date}."
        new_sheet['F3'] = datetime.now().strftime("This report was generated on %A, %B %d, %Y at %I:%M %p.")

        # Creating an Alignment object for center and middle alignment
        align_center_middle = Alignment(horizontal='center', vertical='center', wrap_text=True)

        # Applying the alignment and fill to the cells
        new_sheet['F2'].alignment = align_center_middle
        new_sheet['F2'].fill = light_green_fill
        new_sheet['F3'].alignment = align_center_middle
        new_sheet['F3'].fill = light_green_fill

        # Setting the width of column 'F' to 80 points
        new_sheet.column_dimensions['F'].width = 80

        # Save the new workbook
        today_str = datetime.now().strftime("%Y-%m-%d")
        copy_path = os.path.join(to_sell_folder, f"Products To Sell - {today_str}.xlsx")
        new_workbook.save(copy_path)

        # Open the modified Excel file
        if sys.platform == "win32":
            os.startfile(copy_path)
        elif sys.platform == "darwin":  # macOS
            subprocess.run(["open", copy_path])
        else:  # Linux variants
            subprocess.run(["xdg-open", copy_path])

    def exit_correlate_window(self):
        self.correlate_window.destroy()
        self.Settings_Window_Start()

    def back_to_main(self):
        self.settings_window.destroy()
        self.master.deiconify()
        self.master.state('zoomed')
        
        # Load settings again in case they were changed
        self.load_settings()
        
        # Refresh the folder list with the updated settings
        self.combine_and_display_folders()
        
        self.search_entry.focus_set()

    def choose_inventory_folder(self):
        inventory_folder = filedialog.askdirectory()
        if inventory_folder:
            self.inventory_folder = inventory_folder
            self.inventory_folder_label.config(text=inventory_folder)  # Update the label directly
            self.save_settings()
            self.combine_and_display_folders()

    @staticmethod
    def custom_sort_key(s):
        # A regular expression to match words in the folder name.
        # Words are defined as sequences of alphanumeric characters and underscores.
        words = re.findall(r'\w+', s.lower())
        
        # The key will be a tuple consisting of the length of the first word,
        # the first word itself (for alphanumeric sorting), and then the rest of the words.
        # Lowercase all words for case-insensitive comparison, numbers will sort naturally before letters.
        return (len(words[0]),) + tuple(words)

    def combine_and_display_folders(self):
        # Clear the folder list first
        self.folder_list.delete(0, tk.END)

        # Initialize additional folders based on the inventory folder
        if self.inventory_folder:
            parent_dir = os.path.dirname(self.inventory_folder)
            self.damaged_folder = os.path.join(parent_dir, "Damaged")
            self.personal_folder = os.path.join(parent_dir, "Personal")

            # Create additional folders if they don't exist
            for folder in [self.damaged_folder, self.personal_folder]:
                if not os.path.exists(folder):
                    os.makedirs(folder)

        # Begin a transaction
        self.db_manager.cur.execute("BEGIN")
        try:
            # Combine the folders from all paths including damaged and personal folders
            combined_folders = []
            for folder_path in [self.inventory_folder, self.sold_folder, self.to_sell_folder, self.damaged_folder, self.personal_folder]:
                if folder_path and os.path.exists(folder_path):
                    for root, dirs, files in os.walk(folder_path):
                        for dir_name in dirs:
                            combined_folders.append(dir_name)
                            full_path = os.path.join(root, dir_name)
                            # Update the database with the current folder paths
                            self.db_manager.cur.execute("INSERT OR REPLACE INTO folder_paths (Folder, Path) VALUES (?, ?)", (dir_name, full_path))
            self.db_manager.conn.commit()  # Commit the transaction if all is well
        except Exception as e:
            self.db_manager.conn.rollback()  # Rollback if there was an error
            messagebox.showerror("Database Error", f"An error occurred while updating the folder paths: {e}")
        # Deduplicate folder names
        unique_folders = list(set(combined_folders))

        # Sort using the custom sort key function
        sorted_folders = sorted(unique_folders, key=self.custom_sort_key)

        # Insert the sorted folders into the list widget
        for folder in sorted_folders:
            self.folder_list.insert(tk.END, folder)

    def choose_sold_folder(self):
        self.sold_folder = filedialog.askdirectory()
        if self.sold_folder:
            self.sold_folder_label.config(text=self.sold_folder)  # Update the label directly
            self.save_settings()
        # Update the Sold Folder path
        self.db_manager.cur.execute('''
            INSERT INTO folder_paths (Folder, Path) VALUES ('Sold', ?)
            ON CONFLICT(Folder) DO UPDATE SET Path = excluded.Path;
        ''', (self.sold_folder,))

        self.db_manager.conn.commit()

    def choose_to_sell_folder(self):
        self.to_sell_folder = filedialog.askdirectory()
        if self.to_sell_folder:
            self.to_sell_folder_label.config(text=self.to_sell_folder)
            self.save_settings()  # Save the settings including the new folder path

    def save_settings(self):
        # Here you will gather all the paths and write them to the settings.txt file
        with open("folders_paths.txt", "w") as file:
            file.write(f"{self.inventory_folder}\n{self.sold_folder}\n{self.to_sell_folder}")

    def search(self, event):
        search_terms = self.search_entry.get().split()  # Split the search string into words
        if search_terms:
            self.folder_list.delete(0, tk.END)  # Clear the current list

            # Define a list of folder paths to search in
            search_paths = [
                self.inventory_folder,
                self.sold_folder,
                self.to_sell_folder,
                self.damaged_folder,
                self.personal_folder
            ]

            # Filter out None or invalid paths
            valid_search_paths = [path for path in search_paths if path and os.path.exists(path)]

            # Perform the search in each valid path
            for path in valid_search_paths:
                for root, dirs, files in os.walk(path):
                    # Check if 'dirs' is empty, meaning 'root' is a leaf directory
                    if not dirs:
                        folder_name = os.path.basename(root)  # Get the name of the leaf directory
                        # Check if all search terms are in the folder name (case insensitive)
                        if all(term.upper() in folder_name.upper() for term in search_terms):
                            self.folder_list.insert(tk.END, folder_name)
        else:
            self.combine_and_display_folders()  # If the search box is empty, display all folders        

    def display_product_details(self, event):

        selection = self.folder_list.curselection()
        # Get the index of the selected item
        if not selection:
            return  # No item selected
        index = selection[0]
        selected_folder_name = self.folder_list.get(index)
        selected_product_id = selected_folder_name.split(' ')[0].upper()  # Assuming the product ID is at the beginning
        if self.edit_mode:
            self.toggle_edit_mode()
        # Ensure that the Excel file path and sheet name are set
        filepath, sheet_name = self.load_excel_settings()
        if filepath and sheet_name:
            self.excel_manager.filepath = filepath
            self.excel_manager.sheet_name = sheet_name
            self.excel_manager.load_data()  # Load the data

            # Retrieve product information from the DataFrame
            try:
                product_info = self.excel_manager.get_product_info(selected_product_id)
                # Right after fetching product_info
                self.product_folder_path = self.get_folder_path_from_db(selected_product_id)

                if product_info:
                    
                    self.product_image_label.config(image='')
                    self.product_image_label.configure(text='Loading image...')
                    # 1. Find the column number for "Product Image"
                    product_image_col_num = None
                    for col_num, col_name in enumerate(self.excel_manager.data_frame.columns):
                        if col_name == 'Product Image':
                            product_image_col_num = col_num
                            break
                    # 2. Get the current row number
                    current_row_num = self.excel_manager.data_frame[self.excel_manager.data_frame['Product ID'].str.upper() == selected_product_id.upper()].index[0]
  
                    # 3. Print the column name and row number
                    if product_image_col_num is not None:
                        self.load_and_display_image(current_row_num  +3, product_image_col_num)


                    self.edit_button.config(state="normal")
                    self.order_link_text.config(state='normal')
                    self.cancelled_order_var.set(self.excel_value_to_bool(product_info.get('Cancelled Order')))
                    self.damaged_var.set(self.excel_value_to_bool(product_info.get('Damaged')))
                    self.personal_var.set(self.excel_value_to_bool(product_info.get('Personal')))
                    self.reviewed_var.set(self.excel_value_to_bool(product_info.get('Reviewed')))
                    self.pictures_downloaded_var.set(self.excel_value_to_bool(product_info.get('Pictures Downloaded')))
                    self.sold_var.set(self.excel_value_to_bool(product_info.get('Sold')))
                    
                    # For each field, check if the value is NaN using pd.isnull and set it to an empty string if it is
                    self.asin_var.set('' if pd.isnull(product_info.get('ASIN')) else product_info.get('ASIN', ''))
                    self.product_id_var.set('' if pd.isnull(product_info.get('Product ID')) else product_info.get('Product ID', ''))

                    self.product_name_text.configure(state='normal')
                    self.product_name_text.delete(1.0, "end")
                    product_name = product_info.get('Product Name', '')
                    if product_name:
                        self.product_name_text.insert("insert", product_name) 
                    self.product_name_text.configure(state='disabled')

                    self.comments_text.configure(state='normal')
                    self.comments_text.delete(1.0, "end")
                    comments_text = product_info.get('Comments', '')
                    if comments_text:
                        self.comments_text.insert("insert", comments_text) 
                    self.comments_text.configure(state='disabled')

                    # When a product is selected and the order date is fetched
                    order_date = product_info.get('Order Date', '')
                    formatted_order_date = ''  # Default value
                    if isinstance(order_date, datetime):
                        formatted_order_date = order_date.strftime('%m/%d/%Y')
                        self.order_date_var.set(formatted_order_date)
                    elif isinstance(order_date, str) and order_date:
                        try:
                            # If the date is in the format 'mm/dd/yy', such as '2/15/23'
                            order_date = datetime.strptime(order_date, "%m/%d/%Y")
                            formatted_order_date = order_date.strftime('%m/%d/%Y')
                            self.order_date_var.set(formatted_order_date)
                        except ValueError as e:
                            
                            messagebox.showerror("Error", f"Incorrect date format: {e}")
                    else:
                        self.order_date_var.set('')
                    self.order_date_var.set(formatted_order_date)
                    
                    # When a product is selected and the order date is fetched
                    to_sell_after = product_info.get('To Sell After', '')
                    formatted_to_sell_after = ''  # Default value
                    if pd.notnull(to_sell_after):  # Check if 'To Sell After' is not null
                        try:
                            if isinstance(to_sell_after, datetime):
                                formatted_to_sell_after = to_sell_after.strftime('%m/%d/%Y')
                            elif isinstance(to_sell_after, str) and to_sell_after:
                                to_sell_after = datetime.strptime(to_sell_after, "%m/%d/%Y")
                                formatted_to_sell_after = to_sell_after.strftime('%m/%d/%Y')
                        except ValueError as e:
                            messagebox.showerror("Error", f"Incorrect date format: {e}")
                    self.to_sell_after_var.set(formatted_to_sell_after)
                    self.update_to_sell_after_color()
                    
                    sold_date = product_info.get('Sold Date', '')
                    formatted_sold_date = ''  # Default value

                    if pd.notnull(sold_date):  # Check if 'Sold Date' is not null
                        try:
                            if isinstance(sold_date, datetime):
                                formatted_sold_date = sold_date.strftime('%m/%d/%Y')
                            elif isinstance(sold_date, str) and sold_date:
                                # Parse the date string to a date object and format it
                                sold_date = datetime.strptime(sold_date, "%m/%d/%Y").date()
                                formatted_sold_date = sold_date.strftime('%m/%d/%Y')
                        except ValueError as e:
                            messagebox.showerror("Error", f"Incorrect date format: {e}")
                    else:
                        formatted_sold_date = ''

                    self.sold_date_var.set(formatted_sold_date)

                    def format_price(value):
                        if pd.isnull(value):
                            return ''
                        # Separate the fractional and integer parts
                        fractional, integer = math.modf(value)
                        # If the fractional part is 0, use the integer part; otherwise, format with two decimal places
                        return f"${int(integer) if fractional == 0 else f'{value:.2f}'}"

                    def format_percentage(value):
                        if pd.isnull(value):
                            return ''
                        # Separate the fractional and integer parts
                        fractional, integer = math.modf(value)
                        # If the fractional part is 0, use the integer part; otherwise, format with two decimal places
                        return f"{int(integer) if fractional == 0 else f'{value:.2f}'}%"

                    self.fair_market_value_var.set(format_price(product_info.get('Fair Market Value')))
                    self.discount_var.set(format_price(product_info.get('Discount')))
                    self.percent_discount_var.set(format_percentage(product_info.get('Discount Percentage')))

                    self.regular_product_price_var.set(format_price(product_info.get('Product Price')))
                    self.ivu_tax_var.set(format_price(product_info.get('IVU Tax')))
                    self.product_price_plus_ivu_var.set(format_price(product_info.get('Product Price After IVU')))

                    self.product_price_after_discount_var.set(format_price(product_info.get('Product Price After Discount')))
                    self.ivu_tax_after_discount_var.set(format_price(product_info.get('IVU Tax After Discount')))
                    self.product_price_minus_discount_plus_ivu_var.set(format_price(product_info.get('Product Price After IVU and Discount')))

                    self.sold_price_var.set(format_price(product_info.get('Sold Price')) if not pd.isnull(product_info.get('Sold Price')) else '')

                    self.order_link_text.delete(1.0, "end")
                    hyperlink = product_info.get('Order Link', '')
                    if hyperlink:
                        self.order_link_text.insert("insert", hyperlink, "hyperlink")
                        self.order_link_text.tag_add("hyperlink", "1.0", "end")
                        
                    self.payment_type_var.set('' if pd.isnull(product_info.get('Payment Type')) else product_info.get('Payment Type', ''))
                    # ... continue with other fields as needed ...
                    # Add code here to populate the Sold Date and other date-related fields, if applicable
                    
                    # Fetch the full folder path from the database using the product ID.
                    folder_path = self.get_folder_path_from_db(selected_product_id)

                    # Extract the name of the parent directory (where the product folder is located)
                    parent_folder_name = os.path.basename(os.path.dirname(folder_path)) if folder_path else "No Folder"
                    self.product_folder_var.set(parent_folder_name)

                    # If the folder path exists, update the button to open the product folder when clicked
                    if folder_path and os.path.exists(folder_path):
                        self.product_folder_link.config(command=lambda: self.open_product_folder(folder_path), state='normal')
                    else:
                        self.product_folder_var.set("No Folder")
                        self.product_folder_link.config(state='disabled')
                    
                else:
                    self.edit_button.config(state='disabled')
                    self.cancelled_order_var.set(False)
                    self.damaged_var.set(False)
                    self.personal_var.set(False)
                    self.reviewed_var.set(False)
                    self.pictures_downloaded_var.set(False)
                    self.sold_var.set(False)
                    self.product_image_label.configure(text="Image not loaded.")
                    # Populate the widgets with the matched data
                    self.asin_var.set('')
                    self.product_id_var.set('')
                    self.to_sell_after_var.set('')
                    # Add code here to handle the product image, if applicable
                    self.product_name_text.configure(state='normal')
                    self.product_name_text.delete(1.0, tk.END)
                    self.product_name_text.insert(tk.END, 'Product not found in Excel.')
                    self.product_name_text.configure(state='disabled')
                    self.comments_text.configure(state='normal')
                    self.comments_text.delete(1.0, tk.END)
                    self.comments_text.insert(tk.END, 'Comment not found in Excel.')
                    self.comments_text.configure(state='disabled')
                    self.order_date_var.set('')
                    self.fair_market_value_var.set('')
                    self.discount_var.set('')
                    self.percent_discount_var.set("")

                    self.product_price_after_discount_var.set("")
                    self.ivu_tax_after_discount_var.set("")
                    self.product_price_minus_discount_plus_ivu_var.set("")

                    self.product_price_plus_ivu_var.set('')
                    self.ivu_tax_var.set('')
                    self.regular_product_price_var.set('')
                    self.order_link_text.delete(1.0, "end")
                    self.sold_price_var.set('')
                    self.payment_type_var.set('')
                    self.sold_date_var.set('')
                    self.product_folder_var.set("No Folder")
                    self.product_folder_link.config(state='disabled')
                    self.order_link_text.config(state='disabled')

            except Exception as e:
                messagebox.showerror("Error", f"An error occurred: {e}")
                #print(f"Error retrieving product details: {e}")
        else:
            messagebox.showerror("Error", "Excel file path or sheet name is not set.")

        # Unbind the Enter key from the save_button's command
        self.master.unbind('<Return>')
        self.master.unbind('<Escape>')
       
        # Bind the Enter key to the global enter handler
        self.master.bind('<Return>', self.edit_on_key_handler)

    def load_and_display_image(self, current_row_num, product_image_col_num):
        def task():
            print(f"Starting image loading task: Row {current_row_num}, Column {product_image_col_num}")

            if not self.running:
                print("Task exited: Application no longer running")
                return  # Exit the thread if the application is no longer running

            try:
                wb = openpyxl.load_workbook(self.excel_manager.filepath, data_only=True)
                sheet = wb[self.excel_manager.sheet_name]
                print(f"Workbook loaded: {self.excel_manager.filepath}")

                found_image = False
                for image in sheet._images:
                    print(f"Checking image at Row {image.anchor._from.row}, Column {image.anchor._from.col}")
                    if (image.anchor._from.row == current_row_num and 
                        image.anchor._from.col == product_image_col_num):
                        print("Image found, attempting to load...")
                        image_stream = io.BytesIO(image._data())  # Keep the stream open
                        pil_image = Image.open(image_stream)
                        
                        # Resize the image using PIL
                        desired_size = (100, 100)  # Set the desired size
                        resized_image = pil_image.resize(desired_size)

                        # Convert the resized image to Tkinter PhotoImage
                        tk_photo = ImageTk.PhotoImage(resized_image)

                        if self.running:
                            print("Scheduling image update in main thread")
                            self.after(0, lambda: self.update_image_label(tk_photo))
                        else:
                            print("Skipped image update: Application no longer running")
                        found_image = True
                        break

                if not found_image and self.running:
                    print("No image found at specified cell")
                    self.after(0, lambda: self.product_image_label.config(text="Product image not found"))
                elif not self.running:
                    print("Skipped 'no image found' update: Application no longer running")

                wb.close()
            except Exception as e:
                print(f"Error loading image: {e}")
                if self.running:
                    self.after(0, lambda: self.product_image_label.config(text="Error loading image"))
                else:
                    print("Skipped error message update: Application no longer running")

        threading.Thread(target=task).start()

    def update_image_label(self, tk_photo):
        if self.running:
            print("Updating image label in main thread")
            self.product_image_label.config(image=tk_photo)
            self.product_image_label.image = tk_photo  # Keep a reference
        else:
            print("Skipped updating image label: Application no longer running")

    def open_product_folder(self, folder_path):
        if sys.platform == "win32":
            os.startfile(folder_path)
        elif sys.platform == "darwin":  # macOS
            subprocess.run(["open", folder_path])
        else:  # Linux variants
            subprocess.run(["xdg-open", folder_path])

    def excel_value_to_bool(self, value):
        # Check for NaN explicitly and return False if found
        if pd.isnull(value):
            return False
        if isinstance(value, str):
            return value.strip().lower() in ['yes', 'true', '1']
        elif isinstance(value, (int, float)):
            return bool(value)
        return False

    def update_to_sell_after_color(self):
        # Get today's date
        today = date.today()

        # Get the date from the to_sell_after_var entry
        to_sell_after_str = self.to_sell_after_var.get()
        if to_sell_after_str:
            try:
                # Parse the date string to a date object
                to_sell_after_date = datetime.strptime(to_sell_after_str, "%m/%d/%Y").date()

                # If the to_sell_after date is today or has passed, change the label's background color to green
                if to_sell_after_date <= today:
                    self.to_sell_after_label.config(background='light green')
                else:
                    self.to_sell_after_label.config(background='white')
            except ValueError:
                # If there's a ValueError, it means the string was not in the expected format
                # Handle incorrect date format or clear the background
                self.to_sell_after_label.config(background='white')
    
    def checkbox_click_control(self, var):
        """Controls the checkbox click based on edit mode."""
        if not self.edit_mode:
            # If not in edit mode, prevent changing the checkbox's state
            return "break"  # Stop the event from propagating further
        # If in edit mode, allow the checkbox to change state

    def toggle_edit_mode(self):

        # Toggle the edit mode
        print("toggling edit mode")
        
        # Toggle the edit mode state. Switch self.edit_mode to True if it's currently False, and vice versa. 
        # This line essentially switches between edit and view modes for the application.
        self.edit_mode = not self.edit_mode
        # Set the state variable depending on self.edit_mode. 
        # When in edit mode (self.edit_mode is True), state is set to 'normal', enabling interaction and editing of widgets. 
        # When not in edit mode (self.edit_mode is False), state is set to 'disabled', making widgets non-interactive and uneditable.
        state = 'normal' if self.edit_mode else 'disabled' 
        # Set the readonly_state variable based on self.edit_mode. 
        # Use 'readonly' for specific widgets when in edit mode to allow viewing but restrict modification. 
        # Set to 'disabled' when not in edit mode to make these widgets non-interactive.
        readonly_state = 'readonly' if self.edit_mode else 'disabled'  # Use 'readonly' when in edit mode, 'disabled' otherwise        
        
        self.order_date_entry.config(state='disabled')
        self.sold_date_button.config(state=state)
        self.clear_button.config(state=state)       
        self.to_sell_after_entry.config(state='disabled')
        self.payment_type_combobox.config(state=readonly_state)
        self.asin_entry.config(state=state)
        self.product_id_entry.config(state='disabled')
        self.product_name_text.config(state='disabled')
        self.fair_market_value_entry.config(state='disabled')
        self.regular_product_price_entry.config(state='disabled')
        self.ivu_tax_entry.config(state='disabled')
        self.product_price_plus_ivu_entry.config(state=state)
        self.discount_entry.config(state=state)
        self.percent_discount_entry.config(state=state)
        self.sold_price_entry.config(state=state)
        self.save_button.config(state=state)
        self.comments_text.config(state=state)

        if self.edit_mode:
            self.product_name_text.bind("<Button-1>", lambda e: None)

            # When in edit mode, bind the Enter key to the save_button's command
            self.master.bind('<Return>', self.save_on_key_handler)
            
            # When in edit mode, bind the Escape key to the edit_button's command
            self.master.bind('<Escape>', self.edit_on_key_handler)
        else:
            self.product_name_text.bind("<Button-1>", lambda e: "break")

            # When not in edit mode, unbind the Enter and Escape keys
            self.master.unbind('<Return>')
            self.master.unbind('<Escape>')
            # When in edit mode, bind the Enter key to the save_button's command
            self.master.bind('<Return>', self.edit_on_key_handler)

    def save(self):
        # Extract values from the widgets
        sold_price = self.sold_price_entry.get()
        sold_date = self.sold_date_var.get()  # Assuming it's a StringVar associated with an Entry
        payment_type = self.payment_type_var.get()  # Similarly, for payment type

        # Check if any of the fields have data
        if sold_price or sold_date or payment_type:
            # Check if all required fields are filled
            if not (sold_price and sold_date and payment_type):
                messagebox.showwarning("Incomplete Data", "Please fill in Sold Price, Sold Date, and Payment Type.")
                return  # Return without saving

        # Update the 'Sold' checkbox based on the 'Sold Date' entry
        if self.sold_date_var.get():
            # If 'Sold Date' is not empty, check 'Sold'
            self.sold_var.set(True)
        else:
            # If 'Sold Date' is empty, uncheck 'Sold'
            self.sold_var.set(False)

        def to_float(value):
            try:
                # Remove any non-numeric characters like $ and %, then convert to float
                numeric_value = value.replace('$', '').replace('%', '')
                return float(numeric_value)
            except ValueError:
                # Return the original value if it can't be converted
                return value
            
        def remove_dollar_sign(value):
            return value.replace('$', '') if isinstance(value, str) else value
        
        product_id = self.product_id_var.get().strip().upper()

        # Ensure that the Excel file path and sheet name are set.
        filepath, sheet_name = self.load_excel_settings()

        if not filepath or not sheet_name:
            messagebox.showerror("Error", "Excel file path or sheet name is not set.")
            return

        # Collect the data from the form.
        product_data = {
            'Cancelled Order': self.cancelled_order_var.get(),
            'Damaged': self.damaged_var.get(),
            'Personal': self.personal_var.get(),
            'Reviewed': self.reviewed_var.get(),
            'Pictures Downloaded': self.pictures_downloaded_var.get(),
            'Sold': self.sold_var.get(),
            'To Sell After': self.to_sell_after_var.get(),
            'Product Name': self.product_name_text.get("1.0", tk.END).strip(),
            'Sold Price': self.sold_price_var.get(),
            'Payment Type': self.payment_type_var.get(),
            'Sold Date': self.sold_date_var.get(),
            'Comments': self.comments_text.get("1.0", tk.END).strip(),
            'Fair Market Value': to_float(remove_dollar_sign(self.fair_market_value_var.get())),
            'Discount': to_float(remove_dollar_sign(self.discount_var.get())),
            'Discount Percentage': to_float(remove_dollar_sign(self.percent_discount_var.get())),
            'Product Price': to_float(remove_dollar_sign(self.regular_product_price_var.get())),
            'IVU Tax': to_float(remove_dollar_sign(self.ivu_tax_var.get())),
            'Product Price After IVU': to_float(remove_dollar_sign(self.product_price_plus_ivu_var.get())),

            'Product Price After Discount': to_float(remove_dollar_sign(self.product_price_after_discount_var.get())),
            'IVU Tax After Discount': to_float(remove_dollar_sign(self.ivu_tax_after_discount_var.get())),
            'Product Price After IVU and Discount': to_float(remove_dollar_sign(self.product_price_minus_discount_plus_ivu_var.get())),

            'Sold Price': to_float(remove_dollar_sign(self.sold_price_var.get())),
            # ... and so on for the rest of your form fields.
        }

        # Use the ExcelManager method to save the data.
        try:
            self.excel_manager.save_product_info(product_id, product_data)
            messagebox.showinfo("Success", "Product information updated successfully.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save changes to Excel file: {e}")
            return
        
        # Folder movement logic
        current_folder_path = self.get_folder_path_from_db(product_id)
        if not current_folder_path:
            messagebox.showerror("Error", f"No current folder path found for Product ID {product_id}")
            return
        
        # Initialize variables for folder paths
        damaged_folder_path = os.path.join(os.path.dirname(self.inventory_folder), "Damaged")
        personal_folder_path = os.path.join(os.path.dirname(self.inventory_folder), "Personal")

        # Create Damaged and Personal folders if they do not exist
        for folder in [damaged_folder_path, personal_folder_path]:
            if not os.path.exists(folder):
                os.makedirs(folder)

        # Decide target folder based on checkbox statuses and other conditions
        if self.sold_var.get():
            target_folder_path = self.sold_folder
        elif self.personal_var.get():
            target_folder_path = personal_folder_path
        elif self.damaged_var.get():
            target_folder_path = damaged_folder_path
        else:
            to_sell_after_str = self.to_sell_after_var.get()
            try:
                to_sell_after_date = datetime.strptime(to_sell_after_str, "%m/%d/%Y").date() if to_sell_after_str else None
            except ValueError as e:
                messagebox.showerror("Error", f"Invalid 'To Sell After' date format: {e}")
                return

            today = date.today()
            if to_sell_after_date and to_sell_after_date <= today:
                target_folder_path = self.to_sell_folder
            else:
                target_folder_path = self.inventory_folder
        # Use #print statements to debug the current and target folder paths
        #print(f"Current folder path: {current_folder_path}")
        #print(f"Target folder path: {target_folder_path}")
        # Check if the target folder is determined and it's not the same as the current folder
        if target_folder_path and os.path.isdir(current_folder_path) and current_folder_path != target_folder_path:
            try:
                # Perform the move operation
                new_folder_path = shutil.move(current_folder_path, target_folder_path)
                
                # Save the new folder path in the database
                self.db_manager.save_folder_path(product_id, new_folder_path)
                
                # Refresh the product folder path attribute to the new path
                self.product_folder_path = new_folder_path
                
                # Ensure that changes are committed to the database
                self.db_manager.commit_changes()
                
                messagebox.showinfo("Folder Moved", f"Folder for '{product_id}' moved successfully to the new location.")
                #print(f"Folder for '{product_id}' moved from {current_folder_path} to {new_folder_path}")
                self.refresh_and_select_product(product_id)

            except Exception as e:
                messagebox.showerror("Error", f"Failed to move the folder: {e}")
                self.refresh_and_select_product(product_id)
                
        doc_data = (product_id, product_id, self.product_name_var.get())  # Construct the doc_data tuple
        self.create_word_doc(doc_data, iid="dummy", show_message=True)  # Call create_word_doc with dummy iid

        self.toggle_edit_mode()
        self.search_entry.focus_set()
        # Unbind the Enter and Escape keys
        #self.master.unbind('<Return>')
        #self.master.unbind('<Escape>')
    
        # Additionally, you might want to re-bind the Enter key to the edit_button's command
        # if you want to be able to press Enter to switch to edit mode again
        #self.master.bind('<Return>', lambda e: self.edit_on_key_handler.invoke())

    def refresh_and_select_product(self, product_id):
        # Refresh the list of products
        self.combine_and_display_folders()
        
        # Convert the product_id to uppercase for case-insensitive comparison
        product_id_upper = product_id.upper()

        # Find the index of the product that was just edited
        product_index = None
        for index, product_name in enumerate(self.folder_list.get(0, tk.END)):
            # Use .split() to get the first part of the folder name and compare it in uppercase
            if product_name.split()[0].upper() == product_id_upper:
                product_index = index
                break
        
        # If the product is found in the list, select it
        if product_index is not None:
            self.folder_list.selection_set(product_index)
            self.folder_list.see(product_index)  # Ensure the product is visible in the list
            self.folder_list.event_generate("<<ListboxSelect>>")  # Trigger the event to display product details
        self.toggle_edit_mode()
        self.search_entry.focus_set()

    def get_folder_names_from_db(self):
        self.db_manager.cur.execute("SELECT Folder FROM folder_paths")
        return [row[0] for row in self.db_manager.cur.fetchall()]

    def get_folder_path_from_db(self, product_id):
        # This query assumes that the folder name starts with the product ID followed by a space
        self.db_manager.cur.execute("SELECT Path FROM folder_paths WHERE Folder LIKE ?", (product_id + ' %',))
        result = self.db_manager.cur.fetchone()
        return result[0] if result else None

    def select_excel_database(self):
        filepath = filedialog.askopenfilename(
            title="Select Excel Database",
            filetypes=[("Excel Files", "*.xlsx *.xls *.xlsm")]
        )
        if filepath:
            self.excel_manager.filepath = filepath  # Save the filepath to the ExcelManager instance
            xls = pd.ExcelFile(filepath)
            sheet_names = xls.sheet_names
            if sheet_names:
                # Automatically select the first sheet if available
                self.excel_manager.sheet_name = sheet_names[0]  # Save the sheet name to the ExcelManager instance
                self.save_excel_settings(filepath, sheet_names[0])  # Save settings
                self.excel_manager.load_data()  # Load the data
                self.update_excel_label()  # Update the label
        xls = pd.ExcelFile(filepath)
        sheet_names = xls.sheet_names
        self.ask_sheet_name(sheet_names, filepath)  # Pass filepath here

    def update_excel_label(self):
        excel_db_text = f"{self.excel_manager.filepath} - Sheet: {self.excel_manager.sheet_name}"
        self.excel_db_label.config(text=excel_db_text)

    def ask_sheet_name(self, sheet_names, filepath):
        sheet_window = tk.Toplevel(self)
        sheet_window.title("Select a Sheet")

        listbox = tk.Listbox(sheet_window, exportselection=False)
        listbox.pack(padx=10, pady=10)

        # Populate listbox with sheet names
        for sheet in sheet_names:
            listbox.insert('end', sheet)

        # Set the default selection
        default_sheet_index = sheet_names.index(self.default_sheet) if self.default_sheet in sheet_names else 0
        listbox.selection_set(default_sheet_index)
        listbox.activate(default_sheet_index)

        # Bind double-click event to the listbox
        listbox.bind('<Double-1>', lambda event: self.confirm_sheet_selection(event, listbox, filepath))

        confirm_button = ttk.Button(sheet_window, text="Confirm", command=lambda: self.confirm_sheet_selection(None, listbox, filepath))
        confirm_button.pack(pady=(0, 10))

        sheet_window.wait_window()

    def confirm_sheet_selection(self, event, listbox, filepath):
        selection_index = listbox.curselection()
        if selection_index:
            selected_sheet = listbox.get(selection_index[0])
            self.select_excel_sheet(selected_sheet, filepath)
            listbox.master.destroy()  # Closes the sheet_window

    def select_excel_sheet(self, selected_sheet, filepath):
        # Code to update the ExcelManager with the new sheet and load data
        self.excel_manager.filepath = filepath
        self.excel_manager.sheet_name = selected_sheet
        self.excel_manager.load_data()
        self.update_excel_label()
        self.save_excel_settings(filepath, selected_sheet)

    def save_excel_settings(self, filepath, sheet_name):
        try:
            with open('excel_and_sheet_path.txt', 'w') as f:
                f.write(f"{filepath}\n{sheet_name}")
            self.update_excel_label()  # Update the label when settings are saved
        except Exception as e:
            messagebox.showerror("Error", f"Unable to save settings: {str(e)}")

    def load_excel_settings(self):
        try:
            with open('excel_and_sheet_path.txt', 'r') as f:
                filepath, sheet_name = f.read().strip().split('\n', 1)
                return filepath, sheet_name
        except FileNotFoundError:
            return None, None
        except Exception as e:
            messagebox.showerror("Error", f"Unable to load settings: {str(e)}")
            return None, None

    def correlate_data(self):
        #print("Correlate button pressed")
        
        filepath, sheet_name = self.load_excel_settings()

        # Check if the Excel settings are properly loaded
        if not filepath or not sheet_name:
            messagebox.showerror("Error", "Excel database settings not found.")
            return

        # Load the data into the ExcelManager instance
        self.excel_manager.filepath = filepath  # Set the filepath
        self.excel_manager.sheet_name = sheet_name  # Set the sheet name
        self.excel_manager.load_data()  # Load the data
        
        try:
            # Load Excel data
            df = pd.read_excel(filepath, sheet_name=sheet_name)
            #print("Excel data loaded successfully.")
            product_ids = df['Product ID'].tolist()
            #print(f"Product IDs from Excel: {product_ids}")
        except Exception as e:
            messagebox.showerror("Error", f"Unable to load Excel file: {str(e)}")
            return
            # Filter out nan values from the product_ids list using pandas notnull function
            
        # Sort the DataFrame based on 'Product ID'
        df_sorted = df.sort_values('Product ID').dropna(subset=['Product ID'])

        # Filter out nan values from the product_ids list
        product_ids = df_sorted['Product ID'].tolist()
        #print(f"Sorted and Filtered Product IDs from Excel: {product_ids}")
        
        missing_docs = []
        for product_id in product_ids:
            folder_path = self.get_folder_path_from_db(str(product_id))
            #print(f"Checking folder for Product ID {product_id}: {folder_path}")
            if folder_path:
                word_docs = [f for f in os.listdir(folder_path) if f.endswith('.docx')]
                #print(f"Word documents in folder: {word_docs}")
                if not word_docs:  # If there's no Word document
                    product_name = df.loc[df['Product ID'] == product_id, 'Product Name'].iloc[0]
                    missing_docs.append((os.path.basename(folder_path), product_id, product_name))


        #print(f"Missing documents: {missing_docs}")
        if missing_docs:
            self.prompt_correlation(missing_docs)
        else:
            messagebox.showinfo("Check complete", "No missing Word documents found.")
        # Filter out nan values from the product_ids list

    def update_links_in_excel(self):
        try:
            with open('excel_and_sheet_path.txt', 'r') as file:
                lines = file.readlines()
                excel_path = lines[0].strip()
                sheet_name = lines[1].strip()

            workbook = openpyxl.load_workbook(excel_path)
            sheet = workbook[sheet_name]

            # Find the index of the columns
            header_row = sheet[1]
            product_name_col_index = None
            order_link_col_index = None

            for index, cell in enumerate(header_row):
                if cell.value == 'Product Name':
                    product_name_col_index = index + 1
                elif cell.value == 'Order Link':
                    order_link_col_index = index + 1

            if product_name_col_index is None or order_link_col_index is None:
                messagebox.showerror("Error", "Necessary columns not found.")
                return

            # Iterate through all the rows and update hyperlinks in 'Order Link' column
            for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, max_col=product_name_col_index):
                product_name_cell = row[product_name_col_index - 1]
                order_link_cell = sheet.cell(row=product_name_cell.row, column=order_link_col_index)
                # Add condition here to check if the 'Order Link' cell already has a hyperlink
                if not order_link_cell.hyperlink:  # Only update if the 'Order Link' cell is empty
                    # Copy only the hyperlink URL
                    if product_name_cell.hyperlink:
                        order_link_cell.hyperlink = product_name_cell.hyperlink
                        order_link_cell.value = product_name_cell.hyperlink.target  # Set the cell value to the hyperlink URL

            workbook.save(excel_path)
            messagebox.showinfo("Success", "Links have been updated in the Excel file.")

        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while updating links: {e}")
        
        self.db_manager.delete_all_folders()
        self.db_manager.setup_database()
        self.update_asin_in_excel()

    def update_asin_in_excel(self):
        try:
            with open('excel_and_sheet_path.txt', 'r') as file:
                lines = file.readlines()
                excel_path = lines[0].strip()
                sheet_name = lines[1].strip()

            workbook = openpyxl.load_workbook(excel_path)
            sheet = workbook[sheet_name]

            # Find the index of the columns
            header_row = sheet[1]
            order_link_col_index = None
            asin_col_index = None

            for index, cell in enumerate(header_row):
                if cell.value == 'Order Link':
                    order_link_col_index = index + 1
                elif cell.value == 'ASIN':
                    asin_col_index = index + 1

            if order_link_col_index is None or asin_col_index is None:
                print("Order Link or ASIN columns not found.")  # Debug print
                messagebox.showerror("Error", "Order Link or ASIN columns not found.")
                return

            # Iterate through all the rows and update ASIN based on 'Order Link'
            for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, max_col=order_link_col_index):
                order_link_cell = row[order_link_col_index - 1]
                if order_link_cell.value and '/' in order_link_cell.value:
                    asin_value = order_link_cell.value.split('/')[-1]
                    asin_cell = sheet.cell(row=order_link_cell.row, column=asin_col_index)
                    # Add condition here to check if the ASIN cell is empty
                    if not asin_cell.value:  # Only update if the ASIN cell is empty
                        asin_cell.value = asin_value
                        print(f"Updated ASIN for row {order_link_cell.row}: {asin_value}")  # Debug print

            workbook.save(excel_path)
            print("Excel file saved with updated ASINs.")  # Debug print
            messagebox.showinfo("Success", "ASINs have been updated in the Excel file.")

        except Exception as e:
            print(f"An error occurred while updating ASINs: {e}")  # Debug print
            messagebox.showerror("Error", f"An error occurred while updating ASINs: {e}")
        self.db_manager.delete_all_folders()
        self.db_manager.setup_database()
        self.update_to_sell_after_in_excel()

    def update_to_sell_after_in_excel(self):
        try:
            with open('excel_and_sheet_path.txt', 'r') as file:
                lines = file.readlines()
                excel_path = lines[0].strip()
                sheet_name = lines[1].strip()

            workbook = openpyxl.load_workbook(excel_path)
            sheet = workbook[sheet_name]

            # Find the index of the columns
            header_row = sheet[1]
            order_date_col_index = None
            to_sell_after_col_index = None

            for index, cell in enumerate(header_row):
                if cell.value == 'Order Date':
                    order_date_col_index = index + 1
                elif cell.value == 'To Sell After':
                    to_sell_after_col_index = index + 1

            if order_date_col_index is None or to_sell_after_col_index is None:
                print("Order Date or To Sell After columns not found.")  # Debug print
                messagebox.showerror("Error", "Order Date or To Sell After columns not found.")
                return

            # Iterate through all the rows and update 'To Sell After' based on 'Order Date'
            for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, max_col=order_date_col_index):
                order_date_cell = row[order_date_col_index - 1]
                if order_date_cell.value and isinstance(order_date_cell.value, datetime):
                    to_sell_after_date = order_date_cell.value + relativedelta(months=+6)
                    to_sell_after_cell = sheet.cell(row=order_date_cell.row, column=to_sell_after_col_index)
                    
                    # Add condition here to check if the 'To Sell After' cell is empty
                    if not to_sell_after_cell.value:  # Only update if the 'To Sell After' cell is empty
                        to_sell_after_cell.value = to_sell_after_date
                        print(f"Updated To Sell After for row {order_date_cell.row}: {to_sell_after_date}")  # Debug print

            workbook.save(excel_path)
            print("Excel file saved with updated To Sell After dates.")  # Debug print
            messagebox.showinfo("Success", "To Sell After dates have been updated in the Excel file.")

        except Exception as e:
            print(f"An error occurred while updating To Sell After dates: {e}")  # Debug print
            messagebox.showerror("Error", f"An error occurred while updating To Sell After dates: {e}")

        self.db_manager.delete_all_folders()
        self.db_manager.setup_database()
        self.combine_and_display_folders()

    def update_folder_names(self):
        # Load folder paths from folders_paths.txt
        with open("folders_paths.txt", "r") as file:
            lines = file.read().splitlines()
            self.inventory_folder = lines[0]
            self.sold_folder = lines[1]
            self.to_sell_folder = lines[2] if len(lines) > 2 else None
        
        # Ensure the paths for Damaged and Personal folders are set
        parent_dir = os.path.dirname(self.inventory_folder)
        self.damaged_folder = os.path.join(parent_dir, "Damaged")
        self.personal_folder = os.path.join(parent_dir, "Personal")
        
        # Load Excel path and sheet name from excel_and_sheet_path.txt
        with open('excel_and_sheet_path.txt', 'r') as f:
            excel_path, sheet_name = f.read().strip().split('\n', 1)
        
        # Load Excel data
        df = pd.read_excel(excel_path, sheet_name)

        print("Starting the folder renaming process...")

        # Iterate over each folder in the inventory, sold, and to sell folders
        for folder_path in [self.inventory_folder, self.sold_folder, self.to_sell_folder, self.damaged_folder, self.personal_folder]:
            if folder_path and os.path.exists(folder_path):
                # Instead of comparing folder names directly, create a set for more efficient checks

                for item in os.listdir(folder_path):
                    item_path = os.path.join(folder_path, item)
                    if os.path.isdir(item_path):
                        # Extract the presumed product_id from the folder name
                        presumed_product_id = item.split(' ')[0]

                        # Find the matching product_id in the DataFrame
                        product_info = df[df['Product ID'].str.upper() == presumed_product_id.upper()]
                        if not product_info.empty:
                            # Extract the actual product_id and product_name
                            product_id = product_info['Product ID'].iloc[0]
                            product_name = product_info['Product Name'].iloc[0]
                            
                            # Generate the new folder name and sanitize it
                            new_folder_name = self.replace_invalid_chars(f"{product_id} - {product_name}")
                            new_full_path = self.shorten_path(product_id, product_name, folder_path)
                            new_folder_name_from_path = os.path.basename(new_full_path)

                            # Convert both names to a comparable format
                            comparable_item = self.replace_invalid_chars(item).lower().strip()
                            comparable_new_name = new_folder_name_from_path.lower().strip()

                            # Check if the current folder name is already in the correct format
                            if comparable_item == comparable_new_name:
                                print((f"Folder '{item}' already contains the product name.").encode('utf-8', errors='ignore').decode('cp1252', errors='ignore'))
                                continue

                            # Check if the new full path length is within the limit
                            if len(new_full_path) < 260:
                                try:
                                    os.rename(item_path, new_full_path)
                                    print((f"Renamed '{item}' to '{new_folder_name}'").encode('utf-8', errors='ignore').decode('cp1252', errors='ignore'))
                                except OSError as e:
                                    print((f"Error renaming {item_path} to {new_full_path}: {e}").encode('utf-8', errors='ignore').decode('cp1252', errors='ignore'))
                            else:
                                print((f"Skipped renaming {item_path} due to path length restrictions.").encode('utf-8', errors='ignore').decode('cp1252', errors='ignore'))
                        else:
                            print((f"No matching product info found for folder: {item}").encode('utf-8', errors='ignore').decode('cp1252', errors='ignore'))
        messagebox.showinfo("Done", "Moved and renamed folders")
        self.db_manager.delete_all_folders()
        self.db_manager.setup_database()
        self.combine_and_display_folders()
        
    def replace_invalid_chars(self, filename):
        # Windows filename invalid characters
        invalid_chars = '<>:"/\\|?*'
        for ch in invalid_chars:
            if ch in filename:
                filename = filename.replace(ch, "x")
        return filename

    def shorten_path(self, product_id, product_name, base_path):
        # Windows MAX_PATH is 260 characters
        MAX_PATH = 260
        # Initial maximum length for the product name
        max_name_length = 60

        while max_name_length > 0:
            # Truncate product name to fit
            truncated_product_name = product_name[:max_name_length]

            new_folder_name = f"{product_id} - {truncated_product_name}"
            new_full_path = os.path.join(base_path, new_folder_name)

            # Check if the full path length is within the limit
            if len(new_full_path) <= MAX_PATH:
                return new_full_path
            else:
                # Decrement the maximum name length for the next iteration
                max_name_length -= 1

        # If the loop ends without finding a suitable length, return None or handle appropriately
        print("Unable to shorten the product name sufficiently.")
        return None

    def update_folders_paths(self):
        print("Updating folder paths based on Excel data...")

        # Ensure the Excel file path and sheet name are set
        filepath, sheet_name = self.load_excel_settings()
        if not filepath or not sheet_name:
            print("Excel file path or sheet name is not set.")
            return

        # Load the Excel data
        self.excel_manager.filepath = filepath
        self.excel_manager.sheet_name = sheet_name
        self.excel_manager.load_data()

        # Create and define paths for Damaged and Personal folders
        parent_dir = os.path.dirname(self.inventory_folder)
        damaged_folder = os.path.join(parent_dir, "Damaged")
        personal_folder = os.path.join(parent_dir, "Personal")

        # Ensure Damaged and Personal folders exist
        for folder in [damaged_folder, personal_folder]:
            if not os.path.exists(folder):
                os.makedirs(folder)
                print(f"Created folder: {folder}")

        # Iterate through Inventory folders
        if self.inventory_folder and os.path.exists(self.inventory_folder):
            for root, dirs, _ in os.walk(self.inventory_folder):
                for dir_name in dirs:
                    product_id = dir_name.split(' ')[0]  # Assuming Product ID is the first part of the name
                    product_info = self.excel_manager.get_product_info(product_id)

                    if product_info:
                        sold_status = product_info.get('Sold')
                        damaged_status = product_info.get('Damaged')
                        personal_status = product_info.get('Personal')
                        to_sell_after = product_info.get('To Sell After')

                        if sold_status and isinstance(sold_status, str) and sold_status.upper() == 'YES':
                            self.move_product_folder(root, dir_name, self.sold_folder)
                        elif personal_status and isinstance(personal_status, str) and personal_status.upper() == 'YES':
                            self.move_product_folder(root, dir_name, personal_folder)
                        elif damaged_status and isinstance(damaged_status, str) and damaged_status.upper() == 'YES':
                            self.move_product_folder(root, dir_name, damaged_folder)
                        elif self.is_date_today_or_before(to_sell_after):
                            self.move_product_folder(root, dir_name, self.to_sell_folder)
                        else:
                            print(f"Keeping {dir_name} in Inventory")


        else:
            print(f"Inventory folder not found: {self.inventory_folder}")
        self.db_manager.delete_all_folders()
        self.db_manager.setup_database()
        self.update_folder_names()

    def move_product_folder(self, current_path, folder_name, target_folder):
        if target_folder and os.path.exists(target_folder):
            full_path = os.path.join(current_path, folder_name)

            # Extracting the part of the folder name before the first '-' and keeping the hyphen
            new_folder_name = folder_name.split('-', 1)[0].strip() + ' -'
            new_full_path = os.path.join(target_folder, new_folder_name)

            try:
                # Check if a folder with the new name already exists in the target directory
                if os.path.exists(new_full_path):
                    print(f"Folder with name '{new_folder_name}' already exists in the target directory.")
                    return

                # Rename and move the folder
                os.rename(full_path, new_full_path)
                print(f"Moved and renamed folder '{folder_name}' to '{new_folder_name}' in '{target_folder}'")
            except Exception as e:
                print(f"Error moving folder '{folder_name}': {e}")
        else:
            print(f"Target folder not found: {target_folder}")

    def is_date_today_or_before(self, date_input):
        if pd.isnull(date_input):
            return False

        # Check if the input is already a datetime object
        if isinstance(date_input, datetime):
            to_sell_date = date_input.date()
        else:
            try:
                # If it's a string, parse it into a datetime object
                to_sell_date = datetime.strptime(date_input, "%m/%d/%Y").date()
            except ValueError:
                print(f"Invalid date format: {date_input}")
                return False

        return to_sell_date <= datetime.today().date()

    def rpc_formula(self, fair_market_value):
        tax_rate = Decimal('0.115')
        original_value = (Decimal(fair_market_value) / (1 - tax_rate)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

        # Round up to the nearest 5 or 0
        total_price = -(-original_value // Decimal('5')) * Decimal('5')

        # Calculate the regular product price by subtracting the IVU tax from the total price
        regular_product_price = (total_price / (1 + tax_rate)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

        # Calculate the IVU tax from the total price
        IVU_tax = (regular_product_price * tax_rate).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

        # Calculate the 10% reseller earnings of the regular product price
        price_discount = (regular_product_price * Decimal('0.10')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

        total_price = regular_product_price + IVU_tax

        return regular_product_price, total_price, IVU_tax, price_discount

    def update_prices(self):
        # Read the Excel path and sheet name from the file
        with open('excel_and_sheet_path.txt', 'r') as file:
            excel_path, sheet_name = file.read().strip().split('\n')
        
        # Load the workbook and the specific sheet
        workbook = load_workbook(excel_path)
        sheet = workbook[sheet_name]

        # Convert the sheet into a DataFrame
        data = sheet.values
        columns = next(data)[0:]  # The first row of the sheet contains column names
        df = pd.DataFrame(data, columns=columns)
        df = df[1:]  # Skip the header row

        # Convert columns to 'object' type to avoid FutureWarning
        object_columns = ['Product Price', 'Product Price After IVU', 'IVU Tax', 'Discount']
        for col in object_columns:
            df[col] = df[col].astype('object')

        # Define inner functions for conversions inside update_prices to keep them scoped
        def to_currency(value):
            return "${:,.2f}".format(value)

        def currency_to_float(value):
            if pd.isna(value):
                return 0  # or some other sensible default value
            elif isinstance(value, str) and value.startswith('$'):
                value = value.replace('$', '').replace(',', '')
                try:
                    return float(value)
                except ValueError:
                    return 0  # or some other sensible default value
            return value


        # Iterate through the DataFrame and update the prices
        for index, row in df.iterrows():
            if pd.isna(row['Product Price']) or pd.isna(row['Product Price After IVU']) or pd.isna(row['IVU Tax']):
                fair_market_value_raw = row['Fair Market Value']
                fair_market_value = Decimal(currency_to_float(fair_market_value_raw))
                regular_product_price, total_price, IVU_tax, price_discount = self.rpc_formula(fair_market_value)
                
                # Calculate the discounted prices using Decimal
                product_price_after_discount = (regular_product_price - price_discount).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
                ivu_tax_after_discount = (product_price_after_discount * Decimal('0.115')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
                product_price_plus_ivu_discount = (product_price_after_discount + ivu_tax_after_discount).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

                df.at[index, 'Product Price'] = float(regular_product_price)
                df.at[index, 'Product Price After IVU'] = float(total_price)
                df.at[index, 'IVU Tax'] = float(IVU_tax)
                df.at[index, 'Discount'] = float(price_discount)
                df.at[index, 'Discount Percentage'] = 10  # Assuming a fixed 10% discount
                df.at[index, 'Product Price After Discount'] = float(product_price_after_discount)
                df.at[index, 'IVU Tax After Discount'] = float(ivu_tax_after_discount)
                df.at[index, 'Product Price After IVU and Discount'] = float(product_price_plus_ivu_discount)

        # Clear the existing data in the sheet
        for row in sheet.iter_rows(min_row=2, max_col=sheet.max_column, max_row=sheet.max_row):
            for cell in row:
                cell.value = None

        # Write the updated DataFrame back to the sheet
        for r_idx, df_row in enumerate(dataframe_to_rows(df, index=False, header=False), start=2):
            for c_idx, value in enumerate(df_row, start=1):
                sheet.cell(row=r_idx, column=c_idx, value=value)

        # Save the workbook
        workbook.save(excel_path)
        messagebox.showinfo("Success", "Prices updated successfully in the Excel file.")

    def prompt_correlation(self, missing_docs):
        self.correlate_window = Toplevel(self)
        self.correlate_window.title("Correlate Data")

        self.missing_docs = missing_docs

        # Create a Treeview with columns
        self.correlate_tree = ttk.Treeview(self.correlate_window, columns=('Folder Name', 'Product ID', 'Product Name'), show='headings')
        self.correlate_tree.pack(fill='both', expand=True)

        # Configure the columns
        self.correlate_tree.column('Folder Name', anchor='w', width=150)
        self.correlate_tree.column('Product ID', anchor='center', width=100)
        self.correlate_tree.column('Product Name', anchor='w', width=150)

        # Define the headings
        self.correlate_tree.heading('Folder Name', text='Folder Name', anchor='w')
        self.correlate_tree.heading('Product ID', text='Product ID', anchor='center')
        self.correlate_tree.heading('Product Name', text='Product Name', anchor='w')

        # Add the items to the Treeview
        for i, (folder_name, product_id, product_name) in enumerate(missing_docs):
            self.correlate_tree.insert('', 'end', iid=str(i), values=(folder_name, product_id, product_name))

        # Bind double-click event to an item
        self.correlate_tree.bind('<Double-1>', self.on_item_double_click)
        
        # Adding a Yes to All button
        yes_to_all_button = ttk.Button(self.correlate_window, text="Yes to All", command=self.create_all_word_docs)
        yes_to_all_button.pack()


        exit_button = ttk.Button(self.correlate_window, text="Exit", command=self.exit_correlate_window)
        exit_button.pack()

    def on_item_double_click(self, event):
        item_id = self.correlate_tree.selection()[0]  # Get selected item ID (iid)
        item_values = self.correlate_tree.item(item_id, 'values')
        
        # Convert item values to a doc_data tuple (folder_name, product_id, product_name)
        doc_data = (item_values[0], item_values[1], item_values[2])

        # Call the create_word_doc function with doc_data and the item's iid
        self.create_word_doc(doc_data, item_id)  # show_message is True by default

    def create_all_word_docs(self):
        #print("Create all word docs function called")  # Debug #print statement
        for iid in self.correlate_tree.get_children():
            item_values = self.correlate_tree.item(iid, 'values')
            doc_data = (item_values[0], item_values[1], item_values[2])
            self.create_word_doc(doc_data, iid, show_message=False)
        messagebox.showinfo("Success", "All Word documents have been created.")
        self.correlate_window.destroy()
        self.Settings_Window_Start()

    def create_word_doc(self, doc_data, iid, show_message=True):
        # Unpack the data tuple
        folder_name, product_id, product_name = doc_data
        # Retrieve the folder path from the database
        folder_path = self.get_folder_path_from_db(str(product_id))

        if folder_path:
            try:
                # Retrieve 'To Sell Prices' from the Excel data
                product_price_after_ivu_series = self.excel_manager.data_frame.loc[self.excel_manager.data_frame['Product ID'] == product_id, 'Product Price After IVU']
                if not product_price_after_ivu_series.empty:
                    product_price_after_ivu = product_price_after_ivu_series.iloc[0]
                else:
                    product_price_after_ivu = "N/A"  # Default to "N/A" if not found

                # Retrieve the product link
                order_link_series = self.excel_manager.data_frame.loc[self.excel_manager.data_frame['Product ID'] == product_id, 'Order Link']
                if not order_link_series.empty:
                    order_link = order_link_series.iloc[0]
                else:
                    order_link = "N/A"  # Default to "N/A"            
                    
                    # Retrieve the product 
                comments_series = self.excel_manager.data_frame.loc[self.excel_manager.data_frame['Product ID'] == product_id, 'Comments']
                if not comments_series.empty:
                    comments = comments_series.iloc[0]
                else:
                    comments = "N/A"  # Default to "N/A" 
                    
                    # Retrieve the product 
                product_name_series = self.excel_manager.data_frame.loc[self.excel_manager.data_frame['Product ID'] == product_id, 'Product Name']
                if not product_name_series.empty:
                    product_name = product_name_series.iloc[0]
                else:
                    product_name = "N/A"  # Default to "N/A" 

                    # Retrieve the product 
                discount_series = self.excel_manager.data_frame.loc[self.excel_manager.data_frame['Product ID'] == product_id, 'Discount']
                if not discount_series.empty:
                    discount = discount_series.iloc[0]
                else:
                    discount = "N/A"  # Default to "N/A" 

            except Exception as e:
                print(f"Error retrieving data: {e}")  # Debugging print statement

            # Path for the new Word document
            doc_path = os.path.join(folder_path, f"{product_id}.docx")
            try:
                # Create a new Word document
                doc = Document()
                doc.add_paragraph(f"Product ID: {product_id}")
                doc.add_paragraph(f"Product Name: {product_name}")
                doc.add_paragraph(f"Product Price After IVU: ${product_price_after_ivu}")
                doc.add_paragraph(f"Reseller Earnings: ${discount}")
                doc.add_paragraph(f"Amazon Link(to get the product description and pictures, if needed): {order_link}")
                doc.add_paragraph(f"Comments: {comments}")

                # Save the document
                doc.save(doc_path)

                if show_message:
                    messagebox.showinfo("Document Created", f"Word document for '{product_id}' has been created successfully.")

                # Additional logic (if any)
                if hasattr(self, 'correlate_tree') and not self.correlate_tree.get_children():
                    self.correlate_window.destroy()
                    self.Settings_Window_Start()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to create document for Product ID {product_id}: {e}")
        else:
            messagebox.showerror("Error", f"No folder found for Product ID {product_id}")

    def backup_excel_database(self):
        print("Starting the backup process.")
        # Check if the Excel filepath is set
        if not self.excel_manager.filepath:
            print("No Excel filepath is set.")
            return
        
        # Check if inventory folder exists
        if not self.inventory_folder or not os.path.exists(self.inventory_folder):
            print(f"Inventory folder is not set or does not exist: {self.inventory_folder}")
            return
        
        # Define backup folder, which is alongside the inventory folder
        # We obtain the parent directory of the inventory_folder using os.path.dirname
        parent_dir = os.path.dirname(self.inventory_folder)
        backup_folder = os.path.join(parent_dir, "Excel Backups")
        
        try:
            # Create the backup folder if it doesn't exist
            if not os.path.exists(backup_folder):
                os.makedirs(backup_folder)
                print(f"Backup folder '{backup_folder}' created.")
            else:
                print(f"Backup folder '{backup_folder}' already exists.")
            
            # Generate backup file name
            date_time_str = datetime.now().strftime("%Y-%m-%d - %H-%M-%S")
            backup_filename = f"Backup of {date_time_str}.xlsx"
            backup_path = os.path.join(backup_folder, backup_filename)
            
            # Copy the Excel file to the backup location
            shutil.copy2(self.excel_manager.filepath, backup_path)
            print(f"Backup created at: {backup_path}")
            
            # Double check if the file was actually created
            if not os.path.isfile(backup_path):
                raise FileNotFoundError(f"Backup file not found after copy operation: {backup_path}")
        except Exception as e:
            print(f"Failed to create backup: {e}")
            raise  # Reraise the exception to see the full traceback

    def __del__(self):
        self.db_manager.conn.close()

def exit_application(app, root):
    on_close(app)  # Perform backup
    app.running = False

    root.destroy()  # Exit the application

def main():
    root = ThemedTk(theme="breeze")  # Use any available theme, e.g., "arc"
    root.title("Improved Inventory Manager")
    root.state('zoomed')
    app = Application(master=root)
    
    app.excel_manager.filepath, _ = app.load_excel_settings()


    # Use a lambda to pass 'app' and 'root' to the 'on_close' function
    root.protocol("WM_DELETE_WINDOW", lambda: on_close(app, root))
    
    app.mainloop()

def on_close(app, root):
    
    print("Closing the application and attempting to backup the database.")
    if hasattr(app, 'excel_manager') and app.excel_manager.filepath:
        print(f"Excel file path at time of backup: {app.excel_manager.filepath}")
        try:
            app.backup_excel_database()  # Perform the backup
            print("Backup should now be complete.")
        except Exception as e:
            print(f"An error occurred during backup: {e}")
    else:
        print("Excel manager not set or no filepath available.")
    app.running = False
    root.destroy()  # Call the destroy method to close the application

if __name__ == '__main__':
    main()